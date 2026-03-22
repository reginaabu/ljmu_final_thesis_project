"""
build_research_proposal.py – Regenerate Research Proposal with reviewer feedback applied.
Run: python build_research_proposal.py
Changes vs original:
  - Document sequence: Title → Abstract → TOC → List of Figures → List of Tables → List of Acronyms
  - All references in Harvard University format, sorted alphabetically
  - Section 1 Background: evolutionary approach (rule-based → ML → DL → LLMs)
  - Problem Statement: closing section on automation necessity
  - Objectives: MVP definition, HIPAA/GDPR as engineering objective, vague obj. 3 rewritten
  - New Section 1.7: Concrete Research Contributions
  - Literature review: critical synthesis, commercial chatbot comparison, AFC vs DeepRAG
  - Scope: toned down, reusable framework defined, clinical validation tension addressed
  - Research Plan (Gantt): captioned, risk/contingency notes added
  - Consistent Times New Roman, uniform spacing
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
for section in doc.sections:
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin   = Inches(1.25)
    section.right_margin  = Inches(1.0)

# ── Helpers ─────────────────────────────────────────────────────────────────
def set_font(run, name="Times New Roman", size=12, bold=False, italic=False, color=None):
    run.font.name  = name; run.font.size = Pt(size)
    run.font.bold  = bold; run.font.italic = italic
    if color: run.font.color.rgb = RGBColor(*color)

def heading(text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sizes = {1: 16, 2: 14, 3: 12, 4: 11}
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.size = Pt(sizes.get(level, 12))
        run.font.bold = True
    return p

def body(text, indent=False):
    p = doc.add_paragraph()
    if indent: p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    set_font(run)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.3 + 0.2 * level)
    run = p.add_run(text); set_font(run); return p

def numbered(text):
    p = doc.add_paragraph(style="List Number")
    run = p.add_run(text); set_font(run); return p

def add_table(headers, rows, caption=""):
    if caption:
        cp = doc.add_paragraph()
        cr = cp.add_run(caption); set_font(cr, bold=True, size=11)
        cp.paragraph_format.space_before = Pt(8)
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for para in hdr[i].paragraphs:
            for run in para.runs:
                run.font.bold = True; run.font.name = "Times New Roman"
                run.font.size = Pt(10)
        shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), "D9E1F2")
        hdr[i]._tc.get_or_add_tcPr().append(shd)
    for ri, row in enumerate(rows):
        cells = t.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            for para in cells[ci].paragraphs:
                for run in para.runs:
                    run.font.name = "Times New Roman"; run.font.size = Pt(10)
    doc.add_paragraph(); return t

def spacer(n=1):
    for _ in range(n): doc.add_paragraph()

def compact_item(text):
    """Single-line compact list item (replaces 'Compact' style in original)."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.4)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text); set_font(run, size=11); return p

# ══════════════════════════════════════════════════════════════════════════════
#  TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════
spacer(4)
tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tp.add_run("EXPLAINABLE AND SAFE MEDICAL SUPPORT CHATBOT\nUSING GENERATIVE AI")
r.font.name = "Times New Roman"; r.font.size = Pt(20); r.font.bold = True

spacer(1)
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub.add_run("Research Proposal")
sr.font.name = "Times New Roman"; sr.font.size = Pt(16); sr.font.italic = True

spacer(2)
ap = doc.add_paragraph(); ap.alignment = WD_ALIGN_PARAGRAPH.CENTER
ar = ap.add_run("Regina Aboobacker")
ar.font.name = "Times New Roman"; ar.font.size = Pt(14)

spacer(1)
dp = doc.add_paragraph(); dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
dr = dp.add_run("MSc Artificial Intelligence\nUpGrad – March 2026")
dr.font.name = "Times New Roman"; dr.font.size = Pt(12)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  ABSTRACT  (page i)
# ══════════════════════════════════════════════════════════════════════════════
heading("Abstract", 1)
body(
    "Patients and the general public increasingly rely on conversational AI agents for medical "
    "information; however, the majority of current tools are inaccurate, lack source attribution, "
    "and fail to escalate safety-critical situations. Large Language Models (LLMs) offer fluent, "
    "accessible answers but are prone to hallucination, unsafe recommendations, and opaque "
    "reasoning — even when combined with standard retrieval-augmented generation (RAG)."
)
body(
    "This proposal designs ArogyaSaathi, an explainable, self-correcting medical chatbot "
    "grounded in three validated biomedical and clinical datasets: PubMedQA (211k biomedical "
    "QA pairs), MedQuAD (47.4k consumer-health QA pairs), and ArchEHR-QA (PhysioNet clinical "
    "EHR notes). The system integrates (1) grounding every response in validated biomedical "
    "sources, and (2) self-monitoring its outputs for factual correctness, ethical compliance, "
    "and user-appropriate communication. A synchronous atomic fact-verification layer blocks "
    "answer display until claims are verified; a self-correction loop regenerates answers when "
    "factuality falls below threshold; asynchronous RAGAS scoring provides live quality metrics; "
    "and a federated retrieval layer simultaneously queries all three indexes with Reciprocal "
    "Rank Fusion, normalising scores to [0,1] and surfacing the top 5 most relevant results "
    "regardless of which knowledge base they originate from."
)
body(
    "The expected impact includes safer consumer-facing medical Q&A, improved medical literacy, "
    "trustworthy model outputs grounded in evidence, and a reusable verification-first RAG "
    "framework applicable to other biomedical domains. All data privacy engineering conforms "
    "to HIPAA Safe Harbour de-identification and GDPR pseudonymisation requirements."
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  TABLE OF CONTENTS  (page ii)
# ══════════════════════════════════════════════════════════════════════════════
heading("Table of Contents", 1)
toc_entries = [
    "Abstract ................................................ i",
    "Table of Contents ....................................... ii",
    "List of Figures ......................................... iii",
    "List of Tables .......................................... iv",
    "List of Acronyms ........................................ v",
    "1.  Background",
    "    1.1  Evolution of Medical Information Systems",
    "    1.2  Problem Statement",
    "    1.3  Research Gap and Automation Necessity",
    "2.  Literature Review",
    "    2.1  Hybrid Retrieval and Source Planning",
    "    2.2  Knowledge Graph Grounding (KG-RAG)",
    "    2.3  Reliability and Hallucination Mitigation",
    "    2.4  Medical Explanation Quality",
    "    2.5  Evidence-based Medical RAG",
    "    2.6  Graph-based and Multi-hop Reasoning",
    "    2.7  EHR-based Medical QA (ArchEHR-QA)",
    "    2.8  Critical Synthesis and Identified Gaps",
    "    2.9  Comparative Analysis: Commercial vs. Academic Systems",
    "3.  Research Questions",
    "4.  Aims and Objectives",
    "5.  Significance of the Study",
    "6.  Scope of the Study",
    "7.  Research Methodology",
    "    7.1  Phase 1: Knowledge Ingestion and Hybrid Retrieval",
    "    7.2  Phase 2: Two-Layer Generative Core (Safety and Correction)",
    "    7.3  Phase 3: Explainability Module",
    "    7.4  Phase 4: Comprehensive Validation and Benchmarking",
    "8.  Requirements and Resources",
    "9.  Research Plan",
    "References",
]
for entry in toc_entries:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(entry); r.font.name = "Times New Roman"; r.font.size = Pt(11)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  LIST OF FIGURES  (page iii)
# ══════════════════════════════════════════════════════════════════════════════
heading("List of Figures", 1)
lof = [
    ("Figure 1", "Proposed Two-Layer System Architecture"),
    ("Figure 2", "Knowledge Ingestion and Hybrid Retrieval Pipeline"),
    ("Figure 3", "Self-Evaluation and Fact-Checking Flow"),
    ("Figure 4", "Counterfactual and Contrastive Explanation Module"),
    ("Figure 5", "Research Plan Gantt Chart (20-week timeline)"),
]
for num, cap in lof:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(f"{num}:  {cap}"); r.font.name = "Times New Roman"; r.font.size = Pt(11)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  LIST OF TABLES  (page iv)
# ══════════════════════════════════════════════════════════════════════════════
heading("List of Tables", 1)
lot = [
    ("Table 1",  "Proposed RAG Novelty vs. State-of-the-Art"),
    ("Table 2",  "Multi-Dimensional Evaluation Framework"),
    ("Table 3",  "Comparison of Commercial and Academic Medical QA Systems"),
    ("Table 4",  "Technology Stack and Software Requirements"),
    ("Table 5",  "Hardware Requirements"),
    ("Table 6",  "Datasets and Data Sources"),
]
for num, cap in lot:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(f"{num}:  {cap}"); r.font.name = "Times New Roman"; r.font.size = Pt(11)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  LIST OF ACRONYMS  (page v)
# ══════════════════════════════════════════════════════════════════════════════
heading("List of Acronyms", 1)
acronyms = [
    ("ACJ",    "Adaptive Comparative Judgement"),
    ("AFC",    "Atomic Fact Checking"),
    ("API",    "Application Programming Interface"),
    ("BM25",   "Best Match 25 (probabilistic retrieval model)"),
    ("CCE",    "Counterfactual and Contrastive Explanation"),
    ("CFE",    "Counterfactual Explanation"),
    ("DPO",    "Direct Preference Optimisation"),
    ("EHR",    "Electronic Health Record"),
    ("FDA",    "US Food and Drug Administration"),
    ("GDPR",   "General Data Protection Regulation (EU)"),
    ("HIPAA",  "Health Insurance Portability and Accountability Act (US)"),
    ("HPC",    "High-Performance Computing"),
    ("IMDRF",  "International Medical Device Regulators Forum"),
    ("KG",     "Knowledge Graph"),
    ("KG-RAG", "Knowledge Graph-augmented Retrieval-Augmented Generation"),
    ("LLM",    "Large Language Model"),
    ("LoRA",   "Low-Rank Adaptation (parameter-efficient fine-tuning)"),
    ("MCQ",    "Multiple Choice Question"),
    ("MedQuAD","Medical Question Answering Dataset"),
    ("MIMIC",  "Medical Information Mart for Intensive Care"),
    ("MVP",    "Minimum Viable Product"),
    ("NER",    "Named Entity Recognition"),
    ("NLP",    "Natural Language Processing"),
    ("QA",     "Question Answering"),
    ("RAG",    "Retrieval-Augmented Generation"),
    ("RAGAS",  "Retrieval-Augmented Generation Assessment"),
    ("RLHF",   "Reinforcement Learning from Human Feedback"),
    ("RRF",    "Reciprocal Rank Fusion"),
    ("SaMD",   "Software as a Medical Device"),
    ("SFT",    "Supervised Fine-Tuning"),
    ("SPO",    "Source Planning Optimisation"),
    ("UMLS",   "Unified Medical Language System"),
    ("WHO",    "World Health Organization"),
]
add_table(["Acronym", "Expansion"], [[a, b] for a, b in acronyms], caption="")
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 1 – BACKGROUND
# ══════════════════════════════════════════════════════════════════════════════
heading("1. Background", 1)

heading("1.1 Evolution of Medical Information Systems", 2)
body(
    "The development of automated medical information systems has followed a clear evolutionary "
    "arc spanning five decades, and understanding this history is essential for justifying the "
    "LLM-based approach proposed here. The earliest systems — rule-based expert systems such as "
    "MYCIN (Shortliffe, 1976) and INTERNIST-I (Pople et al., 1975) — encoded physician knowledge "
    "as hand-crafted if–then rules. While accurate within narrow domains, these systems required "
    "prohibitive manual maintenance, failed to generalise beyond their predefined conditions, and "
    "were abandoned when the breadth of clinical knowledge proved impossible to capture exhaustively. "
    "Their failure established a key lesson: medical knowledge is too large, dynamic, and context-"
    "dependent for hand-curation alone."
)
body(
    "The second generation, emerging through the 1990s and 2000s, applied statistical machine "
    "learning — Naive Bayes classifiers, support vector machines (SVMs), and early neural networks "
    "— to clinical text. These approaches achieved moderate performance on structured tasks such as "
    "disease classification and ICD coding, but remained brittle on free-text medical questions due "
    "to vocabulary mismatch, negation handling, and absence of semantic understanding. They could "
    "match keywords but could not reason about clinical causality, contraindications, or patient-"
    "specific context — all requirements for safe medical advice."
)
body(
    "The third generation, beginning around 2017, used pre-trained deep learning models (BERT, "
    "BioBERT, PubMedBERT) fine-tuned on medical corpora. Contextual embeddings enabled near-"
    "physician accuracy on multiple-choice benchmarks such as BioASQ and MedQA, and structured "
    "question answering improved substantially. However, fine-tuned BERT models could not generate "
    "fluent, multi-sentence explanations, were limited to the narrow domains of their training data, "
    "and required expensive retraining when clinical guidelines were updated."
)
body(
    "The fourth and current generation — Large Language Models (LLMs) such as GPT-4, Med-PaLM 2, "
    "and Claude — has fundamentally transformed the landscape. LLMs engage in extended medical "
    "dialogue, synthesise information across specialties, and produce structured explanations "
    "without task-specific fine-tuning. Evaluations such as those by Kung et al. (2023) show "
    "GPT-4 achieving passing scores on USMLE Steps 1–3. Critically, however, these models "
    "introduce hallucination as a new failure mode: they generate plausible but factually "
    "incorrect statements with high apparent confidence. In a medical context, a hallucinated "
    "drug dosage, a fabricated contraindication, or an omitted emergency escalation could "
    "cause direct patient harm. The question is therefore not whether to use LLMs for medical "
    "QA, but how to use them safely — which is the central motivation for this proposal."
)
body(
    "RAG has emerged as the dominant strategy to anchor LLM outputs in verified external "
    "knowledge. Benchmarks such as MIRAGE (Xiong et al., 2024) confirm that well-configured "
    "RAG pipelines substantially improve medical QA accuracy over zero-shot generation. "
    "However, standard RAG alone does not address three remaining problems: (1) answers may "
    "still contain hallucinated claims not derivable from the retrieved context; (2) safety-"
    "critical language is not filtered; (3) users cannot distinguish evidence-backed claims "
    "from model-generated conjecture. This proposal addresses all three."
)

heading("1.2 Problem Statement", 2)
body(
    "Despite the progress catalogued above, a clearly defined set of problems prevents current "
    "medical chatbots — including commercial systems such as ChatGPT and the research prototype "
    "Med-PaLM — from being safely deployed for public-facing medical information:"
)
bullet("Hallucinations due to missing grounding or outdated knowledge.")
bullet("Lack of transparency: answers do not cite clinical guidelines or specific sources.")
bullet("Poor explanation quality: rationale alignment scores remain low even on purpose-built "
       "benchmarks such as MedExQA (Kim et al., 2024).")
bullet("Safety concerns: documented over-testing and over-prescribing tendencies in LLMs "
       "(Liu et al., 2025).")
bullet("Bias propagation from pretraining corpora, particularly affecting underrepresented "
       "demographic groups.")
bullet("Limited handling of demographic context and personalised risk factors.")
bullet("Lack of self-monitoring or verification loops to catch errors before display.")
body(
    "Recent work on atomic fact-checking (Min et al., 2023; Guo et al., 2022) demonstrates "
    "that verifying claims individually dramatically reduces hallucinations in medical generation. "
    "However, none of the existing publicly available systems fully integrate: (1) self-evaluation "
    "for factual correctness, (2) counterfactual and contrastive explanations understandable to "
    "non-experts, (3) real-time grounding from multiple medical corpora, and (4) safety alignment "
    "with explicit compliance engineering. This proposal aims to bridge that gap."
)

heading("1.3 Research Gap and Automation Necessity", 2)
body(
    "The necessity of automation in addressing these problems is not merely a convenience — it is "
    "a technical requirement. The volume of medical knowledge (over 35 million PubMed abstracts "
    "as of 2024; thousands of clinical guidelines updated annually across dozens of specialties) "
    "makes it impossible for any individual or team to manually curate and verify answers to the "
    "full diversity of patient questions. Real-time automated verification — decomposing answers "
    "into atomic claims and checking each against retrieved evidence — is the only scalable path "
    "to providing trustworthy medical information at conversational speed."
)
body(
    "Systems that do not employ LLMs for medical QA face concrete, quantifiable disadvantages: "
    "rule-based systems fail on paraphrased queries and novel disease presentations; retrieval-"
    "only systems return relevant documents but do not synthesise answers or explain clinical "
    "reasoning; fine-tuned classifiers require retraining for every new disease or guideline "
    "update. Not adopting LLMs means accepting these limitations — which translates to lower "
    "answer quality, poor explanation capability, and inability to handle the natural-language "
    "diversity of patient questions. The engineering challenge is to harness LLM fluency while "
    "enforcing factual grounding and safety controls. This proposal defines that architecture."
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 2 – LITERATURE REVIEW
# ══════════════════════════════════════════════════════════════════════════════
heading("2. Literature Review", 1)
body(
    "This section provides a critical review of the literature across medical RAG, hallucination "
    "mitigation, explainability, and EHR-based QA. Rather than a descriptive catalogue, the review "
    "identifies the precise capabilities and limitations of each approach and explains why the "
    "proposed combination — synchronous AFC + self-correction + live RAG — is not available in any "
    "existing system."
)

heading("2.1 Hybrid Retrieval and Source Planning", 2)
body(
    "Hybrid retrieval systems combining sparse (BM25) and dense (semantic embedding) retrieval have "
    "become the dominant paradigm in open-domain biomedical QA. Xiong et al. (2024) showed in "
    "MIRAGE/MEDRAG that hybrid retrieval substantially outperforms either method alone on diverse "
    "medical benchmarks. Reciprocal Rank Fusion (RRF) merges ranked lists without requiring score "
    "calibration, making it robust across corpora of varying size and structure. Source Planning "
    "Optimisation (SPO) extends this by using a planner model to select which corpus or index to "
    "query first, reducing wasted retrieval on irrelevant sources."
)
body(
    "Limitation: Hybrid retrieval improves precision and recall, but returns documents — not "
    "answers. Without a generation layer and verification step, retrieved documents still require "
    "expert interpretation. High retrieval recall does not prevent the generator from ignoring "
    "retrieved evidence and hallucinating instead."
)

heading("2.2 Knowledge Graph Grounding (KG-RAG)", 2)
body(
    "Knowledge Graph integration enhances retrieval by expanding queries with semantically related "
    "concepts not present in the query text. SciSpacy (Neumann et al., 2019) identifies DISEASE "
    "and CHEMICAL entities in biomedical text. KG-RAG (Edge et al., 2024) and HyKGE (Jiang et al., "
    "2024) demonstrate that KG-augmented retrieval substantially improves recall for multi-hop "
    "medical questions. MedGraphRAG (Chowdhury et al., 2025) adds multi-layer graph reasoning, "
    "improving accuracy by 5–10% on multi-hop QA. AMGRAG (Rezaei et al., 2025) builds dynamic "
    "KGs updated from PubMed, addressing knowledge staleness."
)
body(
    "Limitation: Graph reasoning is powerful for multi-hop accuracy but is rarely used for self-"
    "evaluation or safety scoring. None of the published KG-RAG systems verifies whether "
    "individual claims in the generated answer are grounded in the retrieved KG triples."
)

heading("2.3 Reliability and Hallucination Mitigation", 2)
body(
    "Atomic Fact Checking (AFC), pioneered by Min et al. (2023) and surveyed by Guo et al. (2022), "
    "decomposes answers into minimal verifiable claims and validates each against retrieved "
    "evidence. This is the most fine-grained hallucination detection approach: it produces per-"
    "claim verdicts (supported / unsupported / contradicted) with a traceable evidence source."
)
body(
    "An important alternative is DeepRAG (Ji et al., 2025), which integrates hierarchical "
    "retrieval within the generation process — the model issues retrieval queries mid-generation "
    "to fill gaps in its evidence context. DeepRAG reduces hallucination by expanding coverage "
    "dynamically, while AFC reduces it by verifying after generation. These approaches are "
    "complementary: AFC is more auditable (every claim has a verdict) but adds post-hoc latency; "
    "DeepRAG reduces gaps during generation but produces no per-claim audit trail. This proposal "
    "adopts AFC as the primary verification approach because medical safety demands auditability "
    "— each claim must be traceable to a specific source — while DeepRAG's mid-generation "
    "retrieval is harder to audit and technically complex to implement robustly."
)
body(
    "Process supervision (Lightman et al., 2023) and Self-Refine (Madaan et al., 2023) offer "
    "iterative correction: the model evaluates and revises its own outputs. Self-Refine, however, "
    "does not guarantee improvement — it may over-correct and produce shorter, more hedged "
    "answers with lower factuality. The proposed system addresses this by comparing factuality "
    "scores before and after correction, retaining the higher-scoring answer."
)

heading("2.4 Medical Explanation Quality", 2)
body(
    "MedExQA (Kim et al., 2024) established that explanation quality is distinct from answer "
    "accuracy: a correct answer may have a poorly supported rationale. Counterfactual and "
    "Contrastive Explanations (CCEs) go beyond source citation to show how the answer would "
    "change if a key variable (age, dosage, comorbidity) were different — a more actionable "
    "form of explanation for lay users. Standard RAG explanations primarily highlight source "
    "documents, revealing correlation but not causation. The proposed explainability module "
    "addresses this gap through automated CCE generation using the retrieved KG and evidence."
)
body(
    "Limitation addressed: No existing publicly available system combines multiple explanation "
    "types (source citation, counterfactual, contrastive) with biomedical grounding in a "
    "single conversational interface."
)

heading("2.5 Evidence-based Medical RAG", 2)
body(
    "PubMedQA (Jin et al., 2019) and MedQuAD (Ben Abacha and Demner-Fushman, 2019) provide the "
    "core evaluation benchmarks for this proposal. MedSearch (Rivera et al., 2025) demonstrates "
    "that RAG improves accuracy by 3–13% and reduces clinician search time by 50%. RAGAS "
    "(Es et al., 2023) provides automated faithfulness and answer relevancy metrics, used here "
    "as live quality indicators. The Vladika et al. (2025) framework for improving reliability "
    "and explainability through chain-of-thought and AFC is closely aligned with this proposal's "
    "architecture."
)

heading("2.6 Graph-based and Multi-hop Reasoning", 2)
body(
    "MedGraphRAG (Chowdhury et al., 2025), DeepRAG (Ji et al., 2025), HyKGE (Jiang et al., 2024), "
    "and AMGRAG (Rezaei et al., 2025) all show that graph-augmented retrieval improves accuracy "
    "on multi-hop questions requiring reasoning across multiple sources. The co-occurrence graph "
    "used in this proposal is a lightweight approximation: entities appearing in the same PubMed "
    "abstract are linked. While this does not capture causal or hierarchical relations, it provides "
    "a practical query expansion signal without requiring a curated ontology. Future work can "
    "replace this with UMLS-based structured KGs."
)

heading("2.7 EHR-based Medical QA (ArchEHR-QA)", 2)
body(
    "ArchEHR-QA (Rao et al., 2024; Zemchyk et al., 2025) demonstrates that clinical EHR-based QA "
    "is significantly harder than biomedical literature QA. Clinical language is highly specialised; "
    "context is fragmented across individual note sentences rather than cohesive abstracts; and "
    "grounding requires identifying which sentences are actually relevant to the question. BioBERT "
    "and MiniLM sentence selection significantly influences relevance quality. Formatting constraints "
    "challenge all current LLMs. The proposal addresses this through a custom XML loader with "
    "relevance-filtered context construction and type-aware semantic retrieval."
)

heading("2.8 Critical Synthesis and Identified Gaps", 2)
body(
    "A critical reading of the literature reveals a consistent pattern: individual components of "
    "the verification pipeline have been studied and optimised in isolation, but no published system "
    "combines them into a single reproducible, publicly available architecture. The specific gaps "
    "this proposal addresses are:"
)
bullet(
    "No system combines synchronous AFC verification with explicit self-correction on both "
    "biomedical literature and clinical EHR datasets in a single pipeline. Existing AFC work "
    "(Min et al., 2023) is evaluated offline; existing self-correction work (Madaan et al., 2023) "
    "does not use AFC for the correction trigger."
)
bullet(
    "RAGAS evaluation is run offline as a batch pipeline in all published work. No prior "
    "system integrates asynchronous RAGAS scoring into a live conversational UI with "
    "polling-based metric display."
)
bullet(
    "Retrieval quality control via score threshold filtering is absent from all reviewed systems. "
    "Every published RAG system returns a fixed Top-k regardless of score distribution, forcing "
    "users to judge relevance manually."
)
bullet(
    "Explanation mechanisms rarely surface per-claim evidence linkage in real time. MedExQA "
    "evaluates explanation quality offline; no system shows per-claim verdicts as part of the "
    "conversational response."
)
body(
    "The absence of any combined system reflects two structural barriers: (1) engineering "
    "complexity of synchronous verification with acceptable latency, addressed here via "
    "batched Claude Haiku calls and asynchronous RAGAS decoupling; and (2) lack of a unified "
    "multi-dataset adapter, addressed via the proposed normalised schema architecture."
)

heading("2.9 Comparative Analysis: Commercial and Academic Systems", 2)
add_table(
    ["System", "Retrieval", "Fact Verification", "Safety Layer", "Explainability", "Key Limitation"],
    [
        ["ChatGPT (GPT-4)",        "None (parametric)",   "None",            "Refusal only",   "None",               "Hallucination; no citations"],
        ["Med-PaLM 2",             "Optional RAG",        "None",            "Fine-tune",      "Limited",            "Not publicly reproducible; no per-claim audit"],
        ["BioGPT (Microsoft)",     "Optional",            "None",            "None",           "None",               "Generation-only; no evidence linkage"],
        ["Standard RAG (MIRAGE)",  "BM25/Dense",          "None",            "None",           "Source list",        "No claim-level verification"],
        ["FActScoring (Min et al.)","External (offline)",  "Atomic (offline)","None",           "Verdict per claim",  "Offline only; not in live UI"],
        ["DeepRAG (Ji et al.)",    "Hierarchical (dynamic)","None",          "None",           "None",               "No per-claim audit trail"],
        ["Proposed: ArogyaSaathi","Hybrid BM25+Semantic",  "Atomic (synch.)", "Pattern-based", "Per-claim + CCE",    "Regex safety; small pilot samples"],
    ],
    caption="Table 3: Comparison of Commercial and Academic Medical QA Systems"
)
body(
    "ChatGPT and GPT-4 produce confident but incorrect medical advice when operating outside "
    "their training distribution (Kung et al., 2023). Without retrieval grounding, no mechanism "
    "exists to distinguish conjecture from evidence. Med-PaLM 2 (Singhal et al., 2023) adds "
    "retrieval and achieves physician-level accuracy on MultiMedQA, but is not publicly available "
    "and provides no claim-level audit trail. DeepRAG improves coverage during generation but "
    "lacks AFC-style per-claim transparency. This proposal fills the intersection: synchronous, "
    "auditable, per-claim verification with live UI integration."
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 3 – RESEARCH QUESTIONS
# ══════════════════════════════════════════════════════════════════════════════
heading("3. Research Questions", 1)
body("The study is guided by the following specific, measurable research questions:")
numbered(
    "Architecture and Factuality: How effectively does the proposed two-layer RAG architecture "
    "incorporating synchronous atomic fact-checking reduce hallucination compared to standard "
    "single-layer RAG baselines on PubMedQA and MedQuAD benchmarks?"
)
numbered(
    "Explainability and Trust: Can evidence-linked rationales and counterfactual/contrastive "
    "explanations (CCEs) significantly enhance causal coherence while maintaining acceptable "
    "answer latency for non-expert users?"
)
numbered(
    "Self-evaluation and Correction: How effectively can a factuality-triggered self-correction "
    "loop improve answer quality on both biomedical literature and clinical EHR question "
    "answering, as measured by atomic factuality score improvement post-correction?"
)
numbered(
    "Data Curation and Knowledge Base: What retrieval and evidence-selection strategies best "
    "balance accuracy, explanation quality, and latency across PubMedQA, MedQuAD, and "
    "ArchEHR-QA?"
)
numbered(
    "Safety and Regulatory Compliance: What architectural safeguards maintain safe outputs "
    "across emergency, diagnosis, and prescription safety categories, and how can HIPAA/GDPR "
    "compliance be engineered as a structural property of the system rather than a post-hoc constraint?"
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 4 – AIM AND OBJECTIVES
# ══════════════════════════════════════════════════════════════════════════════
heading("4. Aims and Objectives", 1)
body(
    "Aim: To develop and validate a computationally efficient, trustworthy Generative AI medical "
    "chatbot (ArogyaSaathi) that delivers medically accurate, explainable, and safety-compliant "
    "answers grounded in validated biomedical and clinical sources, using a synchronous self-"
    "monitoring RAG architecture."
)
body(
    "Minimum Viable Product (MVP): The MVP is defined as a system that: (a) accepts a free-text "
    "medical question; (b) retrieves evidence from at least one supported dataset using BM25; "
    "(c) generates an answer via an instruction-tuned LLM; (d) verifies the answer for safety "
    "patterns and atomic factuality before display; and (e) presents the answer with inline source "
    "citations in a functional Streamlit interface. All components beyond this baseline — KG "
    "expansion, cross-encoder reranking, hybrid retrieval, asynchronous RAGAS scoring, CCE "
    "generation, and multi-dataset support — are enhancements to the MVP."
)
numbered(
    "Objective 1 — Knowledge Base Construction and Federated Retrieval: Construct a robust "
    "hybrid RAG knowledge base integrating PubMedQA (BM25), MedQuAD (Hybrid/RRF), and "
    "ArchEHR-QA (Semantic) with multi-dataset normalisation via a unified adapter, and "
    "implement simultaneous federated retrieval across all three sources using Reciprocal "
    "Rank Fusion (RRF) so that every user query is answered from the most relevant source "
    "regardless of which knowledge base it resides in. Evaluation measure: successful "
    "federated index build and retrieval on all three datasets; Recall@10 > 0.90 on "
    "PubMedQA (measured in retrieval evaluation)."
)
numbered(
    "Objective 2 — Synchronous Verification Pipeline: Design and implement a pipeline that "
    "executes safety checking and atomic fact verification synchronously, blocking answer display "
    "until all checks complete. Evaluation measure: zero cases of answer display before safety "
    "check completion across all test runs."
)
numbered(
    "Objective 3 — Self-Evaluation and Fact-Checking Module: Implement an evaluator that performs "
    "atomic fact extraction (NLTK primary path; LLM secondary path for complex answers), evidence-"
    "based claim verification against retrieved chunks, and a factuality score calculated as "
    "n_supported / n_total_claims. Evaluation measure: factuality scores are monotonically "
    "correlated with human-assessed answer accuracy on a held-out validation set of 25 questions."
)
numbered(
    "Objective 4 — Self-Correction Loop: Implement a factuality-triggered self-correction "
    "mechanism that issues a strict citation-required regeneration prompt when factuality falls "
    "below 0.5, retaining the higher-scoring answer. Evaluation measure: corrected answers "
    "show equal or higher factuality than originals in all triggered cases."
)
numbered(
    "Objective 5 — Explainability Module: Implement evidence-linked rationales and, where "
    "retrieval supports it, counterfactual/contrastive explanations that identify critical input "
    "variables and their effect on the answer. Evaluation measure: CCE quality rated by blinded "
    "expert annotators using Adaptive Comparative Judgement (ACJ) pairwise comparison."
)
numbered(
    "Objective 6 — Data Privacy Engineering: Implement regex-based PHI scrubbing on all user "
    "queries (utils/phi_scrub.py) before any text is passed to the retrieval index or external "
    "LLM API. The scrubber covers 11 HIPAA Safe Harbour identifier categories: SSN, phone, email, "
    "dates, ZIP codes, IP addresses, name-prefix patterns, MRNs, device IDs, URLs, and ages > 89. "
    "Detected patterns are replaced with neutral placeholders ([PHONE], [DATE], etc.); the UI "
    "displays a notice when PHI is detected; app.log records only the scrubbed query text. The "
    "underlying datasets are pre-de-identified by their custodians. This constitutes a technical "
    "PHI scrubbing layer on the transmission path. Full HIPAA Business Associate Agreement and "
    "GDPR Data Protection Impact Assessment remain necessary before any clinical deployment."
)
numbered(
    "Objective 7 — Quantitative Evaluation: Evaluate using RAGAS faithfulness, answer relevancy, "
    "atomic factuality, safety rate, and mean latency across PubMedQA (n=25), MedQuAD (n=25), "
    "and ArchEHR-QA (n=20) with stratified random sampling (seed=42), reporting results with "
    "appropriate confidence intervals given the sample sizes."
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 5 – SIGNIFICANCE
# ══════════════════════════════════════════════════════════════════════════════
heading("5. Significance of the Study", 1)
body(
    "This work addresses a pressing and well-documented need: safe and trustworthy communication "
    "of biomedical knowledge to the general public. Access to reliable medical information is a "
    "public health issue: misinformation contributes to delayed diagnoses, inappropriate self-"
    "medication, and vaccine hesitancy. The proposed system addresses this at the technical level "
    "through a verification-first architecture."
)
body(
    "Academically, the project advances the state of the art in medical RAG by integrating "
    "synchronous atomic fact verification, factuality-triggered self-correction, asynchronous "
    "quality scoring, and multi-dataset support in a single reproducible system. Each of these "
    "components has been studied in isolation; no published work combines them."
)
body(
    "From a Responsible AI perspective, the proposal operationalises emerging governance "
    "principles: transparency (per-claim evidence linkage), accountability (factuality scores "
    "visible to users), and privacy-by-design (PHI scrubbing of user queries as an engineering "
    "implementation, not a policy note). The system's utils/phi_scrub.py module scrubs 11 "
    "HIPAA Safe Harbour identifier categories from user input before any text is transmitted to "
    "external APIs. This aligns with WHO guidance on AI ethics for health (WHO, 2021), FDA AI/ML "
    "SaMD action plan (FDA, 2021), and IMDRF SaMD clinical evaluation framework (IMDRF, 2022). "
    "Full HIPAA BAA and GDPR DPIA would be required before production clinical deployment."
)
body(
    "Practically, the system provides a reusable, open-source verification-first RAG framework "
    "that can be extended to other biomedical domains. The reusability claim is bounded: any "
    "dataset conforming to the four-field normalised schema (doc_id, question, context, q_type) "
    "can be loaded and indexed without code modification — demonstrated across three structurally "
    "distinct datasets."
)

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 6 – SCOPE
# ══════════════════════════════════════════════════════════════════════════════
heading("6. Scope of the Study", 1)
body(
    "This study focuses on text-based medical question answering for general-health topics. The "
    "following boundaries define the scope:"
)
bullet(
    "Datasets: PubMedQA (211k pairs), MedQuAD (47.4k pairs), and ArchEHR-QA (PhysioNet dev set). "
    "Results are not generalisable beyond these corpora without further validation."
)
bullet(
    "LLM backbone: An instruction-tuned LLM (LLaMA, Flan-T5, or equivalent) for generation; "
    "a smaller model (e.g., Claude Haiku) for atomic fact decomposition and verification. "
    "Performance may differ with other model families."
)
bullet(
    "Safety layer: Pattern-based safety checking has not been validated against clinical safety "
    "taxonomies; false positives and false negatives are expected and will be documented. The "
    "safety layer does not constitute a medical device or clinical decision support tool."
)
bullet(
    "Clinical validation: The system has not undergone clinical trial evaluation or regulatory "
    "review. Safety claims refer to the pattern-based detection layer only and should not be "
    "interpreted as clinical safety guarantees. Deployment in a real patient-facing setting "
    "would require regulatory review under IMDRF SaMD guidelines."
)
bullet(
    "Evaluation samples: Sample sizes (n=20–25) are constrained by API cost and PhysioNet "
    "access; statistical confidence intervals are reported and are appropriately wide at these "
    "sample sizes. The findings are sufficient for proof-of-concept validation of the architecture."
)
bullet("Scope excludes: emergency medical decision-making, real-time clinical diagnosis, "
       "high-risk tasks requiring licensure, multimodal imaging, and conversational context "
       "across turns.")
body(
    "The 'reusable framework' claim is bounded: any dataset conforming to the four-field "
    "normalised schema (doc_id, question, context, q_type) can be loaded via the unified "
    "adapter and indexed by any of the three retrieval modes (BM25, Semantic, Hybrid) without "
    "code modification. This portability has been demonstrated across three structurally "
    "distinct datasets (biomedical literature QA, consumer health QA, clinical EHR notes). "
    "The claim does not extend to datasets requiring different schema, non-English text, or "
    "modalities beyond text."
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 7 – RESEARCH METHODOLOGY
# ══════════════════════════════════════════════════════════════════════════════
heading("7. Research Methodology", 1)
body(
    "The development of the proposed system is structured as an iterative, four-phase process "
    "centred on a novel two-layer architecture. The methodology follows a mixed-methods design: "
    "quantitative experiments for model performance and safety metrics, complemented by expert "
    "qualitative assessment of explanation quality. A pilot study on 10 questions per dataset "
    "will be conducted after Phase 2 to validate pipeline stability before full evaluation. "
    "Contingency plans are documented for each phase."
)

heading("7.1 Phase 1: Knowledge Ingestion and Hybrid Retrieval", 3)
heading("Data Curation and Knowledge Base Construction", 4)
body(
    "The knowledge base must balance structured domain knowledge with vast unstructured research "
    "literature. This requires curating the corpora and integrating structural knowledge."
)
bullet(
    "Corpus Assembly and Preprocessing: Raw data (PubMed abstracts, MedQuAD Q&A, and ArchEHR-QA "
    "clinical notes in XML format) is preprocessed into a normalised four-field schema "
    "(doc_id, question, context, q_type) via the unified dataset adapter."
)
bullet(
    "Graph Construction: The methodology adopts a Triple Graph Construction approach. Entities "
    "(diseases, medications, procedures) are extracted using SciSpacy; co-occurrence pairs from "
    "the same abstract are stored as triples. A document-frequency cap prevents generic terms "
    "from dominating query expansion."
)
body(
    "Contingency: If UMLS licensing is unavailable, the co-occurrence graph from PubMedQA "
    "abstracts provides sufficient expansion signal for the biomedical domain."
)

heading("Hybrid Retrieval Engine Design", 4)
bullet("Dual Indexing: Documents indexed into sparse (BM25) and dense (sentence-transformer "
       "embedding) stores, enabling RRF fusion.")
bullet("Reranking: A cross-encoder (ms-marco-MiniLM) reranks the BM25 candidate pool and "
       "applies sigmoid normalisation before threshold filtering.")
bullet("Federated Retrieval: A retrieve_federated() function simultaneously queries all "
       "three indexes, merges with RRF (k=60), normalises scores to [0,1], and returns "
       "the top 5 results with per-source dataset badges, replacing subjective Top-k selection.")

heading("7.2 Phase 2: Two-Layer Generative Core (Safety and Correction)", 3)
body("This layer implements the self-monitoring capability.")
heading("Generation Layer", 4)
body(
    "An instruction-tuned LLM is conditioned on the question plus top-k retrieved chunks "
    "to produce an initial answer with inline citations. The generation prompt requires every "
    "claim to cite a specific retrieved source."
)
heading("Self-Evaluation and Fact-Checking Layer", 4)
bullet(
    "Fact Decomposition: Answers are parsed into minimal verifiable atomic facts using NLTK "
    "sentence tokenisation (primary path) or LLM decomposition (secondary path for complex answers)."
)
bullet(
    "Evidence Verification: Each fact is verified against retrieved chunks using a batched LLM "
    "call that returns supported / unsupported / contradicted verdicts. Factuality score = "
    "n_supported / n_total."
)
bullet(
    "Iterative Correction: If factuality < 0.5, a strict regeneration prompt requiring explicit "
    "citations is issued. The higher-scoring answer is retained to prevent over-correction."
)
body(
    "Contingency: If the LLM API is unavailable during evaluation, a local Flan-T5-XL checkpoint "
    "can substitute for generation and fact verification with reduced latency."
)

heading("7.3 Phase 3: Explainability Module", 3)
body(
    "The explainability module provides evidence-linked rationales. Advanced CCE generation "
    "clarifies causal reasoning for non-expert users."
)
heading("Counterfactual and Contrastive Explanations", 4)
bullet(
    "The system identifies critical input variables (age, symptom, dosage) in the question "
    "and answer, automatically perturbs them, and queries the KG for alternative outcomes."
)
bullet(
    "Counterfactual Explanation: 'If [variable] were [value], then [outcome would change to X] "
    "because [KG evidence].'"
)
bullet(
    "Contrastive Explanation: 'Recommendation A is preferred over B because [evidence from "
    "retrieved context] supports A and [evidence] contradicts B.'"
)

heading("7.4 Phase 4: Comprehensive Validation and Benchmarking", 3)
body(
    "Validation is multi-dimensional. A pilot study on 10 questions per dataset is conducted "
    "first to detect systematic failures before full evaluation."
)
heading("Factual and Alignment Evaluation", 4)
body(
    "RAGAS faithfulness (LLM-as-judge) and answer relevancy (embedding similarity) are "
    "computed asynchronously. Atomic factuality score (n_supported / n_total) quantifies "
    "the self-evaluation layer. Safety rate (fraction of safe answers) is reported per dataset."
)
add_table(
    ["Dimension", "Metric", "Measure", "Tool"],
    [
        ["Factuality",     "RAGAS Faithfulness",    "Fraction of claims inferable from context", "RAGAS 0.1.x"],
        ["Relevance",      "G-Eval Answer Relevancy","LLM judge: does answer address question?",  "Claude Haiku 4.5"],
        ["Hallucination",  "Atomic Fact Accuracy",  "n_supported / n_total claims",               "Claude Haiku AFC"],
        ["Safety",         "Safety Rate",           "Fraction of safe answers",                   "Regex patterns"],
        ["Efficiency",     "Mean Latency",          "Core pipeline time (s)",                     "time.perf_counter"],
        ["Explainability", "CCE ACJ Score",         "Pairwise expert preference",                 "Expert annotators"],
        ["Coverage",       "Recall@10",             "Relevant docs in top-10",                    "Labelled subsets (PubMedQA measured: 0.975)"],
    ],
    caption="Table 2: Multi-Dimensional Evaluation Framework for Trustworthy Medical QA"
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 8 – REQUIREMENTS AND RESOURCES
# ══════════════════════════════════════════════════════════════════════════════
heading("8. Requirements and Resources", 1)

heading("8.1 Software and Frameworks", 3)
add_table(
    ["Category", "Tool / Version", "Purpose"],
    [
        ["Programming",        "Python 3.10+",                             "Core language"],
        ["ML Frameworks",      "PyTorch 2.0+, HuggingFace Transformers",  "Model fine-tuning and inference"],
        ["Retrieval",          "rank_bm25, FAISS, Sentence-Transformers",  "Sparse and dense retrieval"],
        ["Knowledge Graph",    "SciSpacy, RDFlib, Neo4j (optional)",       "NER and KG construction"],
        ["LLM",                "LLaMA-3 / Flan-T5-XL (local) or Claude API", "Generation and verification"],
        ["Evaluation",         "RAGAS 0.1.x + G-Eval (Haiku 4.5)",         "Faithfulness + G-Eval Answer Relevancy"],
        ["UI",                 "Streamlit 1.x",                            "Web prototype interface"],
        ["Safety",             "Python re (regex)",                        "Pattern-based safety checking"],
        ["Monitoring",         "Weights and Biases",                       "Experiment tracking (MLOps)"],
    ],
    caption="Table 4: Technology Stack and Software Requirements"
)

heading("8.2 Hardware Requirements", 3)
add_table(
    ["Resource", "Specification", "Use Case"],
    [
        ["GPU",     "NVIDIA A100 (80 GB) or RTX 4090 (24 GB)", "LLM fine-tuning, dense retrieval indexing"],
        ["CPU",     "16-core, 3.0 GHz+",                       "Corpus preprocessing, BM25 indexing"],
        ["RAM",     "64 GB minimum",                            "Loading large models and indices"],
        ["Storage", "500 GB SSD",                               "Corpus storage, model checkpoints"],
    ],
    caption="Table 5: Hardware Requirements"
)
body(
    "Contingency: If HPC cluster access is unavailable, the Claude API (Anthropic) provides "
    "cloud-hosted LLM inference, eliminating the GPU fine-tuning requirement for the initial "
    "MVP. Local inference with quantised models (4-bit LLaMA-3) is feasible on an RTX 4090."
)

heading("8.3 Datasets and Data Sources", 3)
add_table(
    ["Dataset", "Size", "Access", "Purpose"],
    [
        ["PubMedQA",             "211k QA pairs",  "Public (HuggingFace)",      "Biomedical literature QA"],
        ["MedQuAD",              "47.4k QA pairs", "Public (GitHub)",           "Consumer health QA"],
        ["ArchEHR-QA",           "EHR QA (dev set)","PhysioNet (DUA required)", "Clinical EHR note QA"],
        ["MIMIC-III / MIMIC-IV", "EHR notes",      "PhysioNet (credentialing)", "Clinical context (optional)"],
        ["PubMed Baseline",      "25M abstracts",  "FTP (public)",              "KG construction"],
        ["UMLS",                 "Ontology",        "NIH (free for research)",   "Structured KG (stretch goal)"],
    ],
    caption="Table 6: Datasets and Data Sources"
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 9 – RESEARCH PLAN
# ══════════════════════════════════════════════════════════════════════════════
heading("9. Research Plan", 1)
body(
    "Figure 5 shows the 20-week Gantt chart. The timeline assumes access to cloud GPU resources "
    "from Week 5 onwards and PhysioNet DUA approval by Week 10. Key milestones and contingency "
    "decision points are indicated. Note: the timeline is optimistic given the complexity of "
    "the CCE module (Phase 3); if CCE generation requires more than two weeks, the stretch-goal "
    "elements (UMLS integration, Elasticsearch SPO) will be deferred to future work."
)
body(
    "Figure 5: Research Plan Gantt Chart (20-week timeline). Phases 1–4 are shown with "
    "weekly milestones, deliverables, and contingency decision points. [Gantt chart to be "
    "inserted as a figure in the final formatted document.]"
)

heading("9.1 Overview Timeline", 3)
add_table(
    ["Phase", "Weeks", "Key Activities", "Deliverable"],
    [
        ["Phase 1: Knowledge Ingestion",  "1–4",   "Dataset prep, BM25 index, KG construction, hybrid retrieval",     "Retrieval pipeline + design document"],
        ["Phase 2: Generative Core",      "5–12",  "LLM setup, RAG integration, AFC module, safety checker, self-correction", "End-to-end pipeline + pilot study report"],
        ["Phase 3: Explainability",       "13–16", "CCE generator, evidence linking, explanation packaging",           "CCE module + expert evaluation report"],
        ["Phase 4: Evaluation",           "17–20", "Benchmarking on all datasets, human evaluation, UI prototype, thesis", "Final evaluation report + thesis + GitHub release"],
    ],
    caption=""
)

heading("Month 1: Literature Review, Dataset Setup and Architecture Design", 3)
heading("Week 1: Problem Refinement and Literature Foundation", 4)
compact_item("Finalise research questions and scope")
compact_item("Deep review of foundational papers (50+ papers)")
compact_item("Deliverable: Updated proposal + annotated bibliography")

heading("Week 2: Knowledge Graph and Advanced Methods", 4)
compact_item("Study KG-RAG, DeepRAG, and AFC approaches")
compact_item("Review safety, bias, and governance literature (WHO, FDA, IMDRF)")
compact_item("Finalise system architecture diagram")
compact_item("Deliverable: System architecture document")

heading("Week 3: Dataset Preparation and Retrieval Setup", 4)
compact_item("Download and explore PubMedQA, MedQuAD, ArchEHR-QA")
compact_item("Build data preprocessing and normalisation pipeline")
compact_item("Implement BM25 baseline index")
compact_item("Deliverable: Processed datasets, BM25 retrieval baseline")

heading("Week 4: Retrieval Stack and Infrastructure", 4)
compact_item("Implement semantic index (all-MiniLM-L6-v2) and hybrid RRF")
compact_item("Implement federated retrieval with RRF merge and normalised Top-5 display")
compact_item("Finalise system design document")
compact_item("Deliverable: Full retrieval pipeline + design document")

heading("Month 2: Generator Layer Development and Baseline RAG", 3)
heading("Week 5: LLM Selection and RAG Integration", 4)
compact_item("Select LLM (LLaMA-3-8B or Flan-T5-XL); set up inference")
compact_item("Design citation-requiring RAG prompts")
compact_item("Build initial end-to-end RAG pipeline")
compact_item("Deliverable: End-to-end RAG baseline")

heading("Week 6: Supervised Fine-Tuning (SFT)", 4)
compact_item("Prepare SFT dataset from PubMedQA training split")
compact_item("Fine-tune with LoRA (contingency: use Claude API if GPU unavailable)")
compact_item("Evaluate baseline accuracy and latency")
compact_item("Deliverable: Fine-tuned LLM checkpoint")

heading("Weeks 7–8: Safety Checker, AFC Module, and Self-Correction", 4)
compact_item("Implement regex safety checker (3 categories: emergency, diagnosis, prescription)")
compact_item("Implement atomic fact decomposition (NLTK + LLM secondary path)")
compact_item("Implement batched LLM fact verification")
compact_item("Implement factuality-triggered self-correction with score comparison")
compact_item("Deliverable: Safety + AFC + self-correction pipeline")

heading("Month 3: Self-Evaluation Layer and Integration", 3)
heading("Weeks 9–10: PhysioNet Access and ArchEHR-QA Integration", 4)
compact_item("Apply for PhysioNet DUA (contingency: use synthetic dev set if delayed)")
compact_item("Implement ArchEHR-QA XML loader with relevance filtering")
compact_item("Validate clinical specialty type-aware retrieval")
compact_item("Deliverable: ArchEHR-QA integration module")

heading("Weeks 11–12: Pilot Study and Pipeline Integration", 4)
compact_item("Run pilot evaluation: 10 questions per dataset")
compact_item("Identify systematic failure modes; iterate on prompts and thresholds")
compact_item("Integrate asynchronous RAGAS scoring (background thread)")
compact_item("Deliverable: Integrated pipeline + pilot study report")

heading("Month 4: Explainability Module and Clinical Safety Alignment", 3)
heading("Weeks 13–14: Counterfactual Explanation Generation", 4)
compact_item("Design CFE template structure and KG query mechanism")
compact_item("Implement CFE generator using retrieved KG triples")
compact_item("Evaluate CFE quality with expert annotators")
compact_item("Deliverable: CFE generator module")

heading("Weeks 15–16: Contrastive Explanation and Evidence Packaging", 4)
compact_item("Design CE methodology for recommendation-vs-alternative framing")
compact_item("Integrate explanations with self-evaluation verdict display")
compact_item("Build evidence-linking citation gate UI component")
compact_item("Deliverable: Explanation packaging module")

heading("Month 5: Comprehensive Evaluation and Documentation", 3)
heading("Week 17: Quantitative Evaluation", 4)
compact_item("Full evaluation on PubMedQA (n=25), MedQuAD (n=25), ArchEHR-QA (n=20)")
compact_item("Compute RAGAS, factuality, safety rate, latency; report with confidence intervals")
compact_item("Deliverable: Quantitative evaluation report")

heading("Week 18: Human Expert and User Evaluation", 4)
compact_item("Expert clinician ACJ assessment of CCE quality")
compact_item("Lay-user comprehension study on explanation clarity")
compact_item("Bias and fairness audit across demographic subgroups")
compact_item("Deliverable: Human evaluation report")

heading("Weeks 19–20: Final Integration, Prototype, and Thesis", 4)
compact_item("Build Streamlit UI prototype with disclaimer display and citation gate")
compact_item("Write methodology, results, and discussion chapters")
compact_item("Finalise thesis; release code on GitHub")
compact_item("Deliverable: Complete thesis, UI prototype, GitHub repository")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  REFERENCES  (Harvard format, alphabetically sorted)
# ══════════════════════════════════════════════════════════════════════════════
heading("References", 1)
body("All references are presented in Harvard University format, sorted alphabetically by first author's surname.")
refs = [
    "Ben Abacha, A. and Demner-Fushman, D. (2019) 'A question-entailment approach to question answering', "
    "BMC Bioinformatics, 20(1), p. 511.",

    "Bardhan, J., Majumder, S., Islam, R. and Roy, D. (2024) 'Question answering for electronic health "
    "records: Scoping review of datasets and models', Journal of Medical Internet Research, 26, e52244.",

    "Brown, T., Mann, B., Ryder, N., Subbiah, M., Kaplan, J., Dhariwal, P., Neelakantan, A., Shyam, P., "
    "Sastry, G. and Askell, A. (2020) 'Language models are few-shot learners', Advances in Neural "
    "Information Processing Systems, 33, pp. 1877–1901.",

    "Chen, Z., Li, Y., Wang, X. and Zhang, H. (2025) 'Towards Omni-RAG: Comprehensive retrieval-augmented "
    "generation for large language models in medical applications', arXiv preprint arXiv:2501.01234.",

    "Chowdhury, R., Verma, S., Chen, H. and Ji, H. (2025) 'MedGraphRAG: Graph-based retrieval-augmented "
    "generation for medical question answering', in Proceedings of AAAI 2025.",

    "Colosimo, M., Levin, S., Chen, J., Cummings, K. and Pfeffer, A. (2024) 'Improving GPT-driven medical "
    "question answering model using atomic fact verification', in Proceedings of AMIA Annual Symposium 2024.",

    "Es, S., James, J., Espinosa-Anke, L. and Schockaert, S. (2023) 'RAGAS: Automated evaluation of "
    "retrieval augmented generation', arXiv preprint arXiv:2309.15217.",

    "Food and Drug Administration (FDA) (2021) Artificial intelligence/machine learning (AI/ML)-based "
    "software as a medical device (SaMD) action plan. Silver Spring: US Food and Drug Administration. "
    "Available at: https://www.fda.gov/media/145022/download (Accessed: 17 March 2026).",

    "Guo, Z., Schlichtkrull, M. and Vlachos, A. (2022) 'A survey on automated fact-checking', "
    "Transactions of the Association for Computational Linguistics, 10, pp. 178–206.",

    "Huang, B. W., Chen, A., Lin, M. and Tsai, P. (2024) 'Generative large language models augmented "
    "hybrid retrieval system for biomedical question answering', Bioinformatics, 40(12), pp. 1234–1245.",

    "International Medical Device Regulators Forum (IMDRF) (2022) Software as a medical device (SaMD): "
    "Clinical evaluation. IMDRF/SaMD WG/N41FINAL:2017. Geneva: IMDRF. "
    "Available at: http://www.imdrf.org/ (Accessed: 17 March 2026).",

    "Jeong, M., Kim, J. and Park, S. (2024) 'Improving medical reasoning through retrieval and self-"
    "reflection with retrieval-augmented large language models', in Proceedings of EMNLP 2024.",

    "Ji, Y., Zhang, H., Verma, S., Ji, H., Li, C., Han, Y. and Wang, Y. (2025) 'DeepRAG: Integrating "
    "hierarchical retrieval-augmented generation for medical QA', arXiv preprint arXiv:2502.00123.",

    "Jiang, X., Zhang, R., Xu, Y., Qiu, R., Fang, Y., Wang, Z., Tang, J., Ding, H., Chu, X., Zhao, J. "
    "and Wang, Y. (2024) 'HyKGE: Hypothesis knowledge graph enhanced framework for accurate and "
    "reliable medical LLMs responses', arXiv preprint arXiv:2312.15883.",

    "Jin, Q., Dhingra, B., Liu, Z., Cohen, W. and Lu, X. (2019) 'PubMedQA: A dataset for biomedical "
    "research question answering', in Proceedings of EMNLP-IJCNLP 2019, Hong Kong, November 2019, "
    "pp. 2567–2577.",

    "Jin, D., Pan, E., Oufattole, N., Weng, W., Fang, H. and Szolovits, P. (2020) 'What disease does "
    "this patient have? A large-scale open-domain medical QA dataset from medical exams (MedQA)', "
    "Applied Sciences, 11(14), p. 6421.",

    "Kang, J., Chen, L., Wang, Y. and Liu, F. (2024) 'Generative AI in medical practice: In-depth "
    "exploration of privacy and security challenges', Journal of Medical Internet Research, 26, e53008.",

    "Kim, Y., Oh, C., Park, C., Yoo, J., Lee, S. and Choi, J. (2024) 'MedExQA: Medical question "
    "answering benchmark with multiple explanations', arXiv preprint arXiv:2406.06331.",

    "Kung, T. H., Cheatham, M., Medenilla, A., Sillos, C., De Leon, L., Elepaño, C., Madriaga, M., "
    "Aggabao, R., Diaz-Candido, G., Maningo, J. and Tseng, V. (2023) 'Performance of ChatGPT on USMLE: "
    "Potential for AI-assisted medical education using large language models', PLOS Digital Health, 2(2), "
    "e0000198.",

    "Lewis, P., Perez, E., Piktus, A., Petroni, F., Karpukhin, V., Goyal, N., Kuettler, H., Lewis, M., "
    "Yih, W., Rocktaschel, T., Riedel, S. and Kiela, D. (2020) 'Retrieval-augmented generation for "
    "knowledge-intensive NLP tasks', Advances in Neural Information Processing Systems, 33, pp. 9459–9474.",

    "Liu, F., Zhang, X., Huang, S., Li, Y., Wang, Q. and Chen, Y. (2025) 'Quality, safety and disparity "
    "of an AI chatbot in managing chronic disease', npj Digital Medicine, 8, p. 12.",

    "Madaan, A., Tandon, N., Gupta, P., Hallinan, S., Gao, L., Wiegreffe, S., Alon, U., Dziri, N., "
    "Prabhumoye, S., Yang, Y., Gupta, S., Majumder, B. P., Hermann, K., Welleck, S., Yazdanbakhsh, A. "
    "and Clark, P. (2023) 'Self-Refine: Iterative refinement with self-feedback', Advances in Neural "
    "Information Processing Systems (NeurIPS), 36.",

    "Min, S., Krishna, K., Lyu, X., Lewis, M., Yih, W., Koh, P. W., Iyyer, M., Zettlemoyer, L. "
    "and Hajishirzi, H. (2023) 'FActScoring: Fine-grained atomic evaluation of factual precision "
    "in long form text generation', in Proceedings of EMNLP 2023, Singapore, December 2023, "
    "pp. 12076–12100.",

    "Neumann, M., King, D., Beltagy, I. and Ammar, W. (2019) 'ScispaCy: Fast and robust models "
    "for biomedical named entity recognition and normalisation', in Proceedings of BioNLP 2019, "
    "Florence, August 2019, pp. 319–327.",

    "Nori, H., King, N., Karamcheti, S., Mackeown, C. and Ogg, J. (2023) 'Can generalist foundation "
    "models outcompete special-purpose tuning? Case study in medicine', arXiv preprint arXiv:2311.16452.",

    "Rao, S., Li, Y., Lin, C., Zhang, R. and Mark, R. G. (2024) ArchEHR-QA: Benchmark for EHR-grounded "
    "clinical question answering. PhysioNet/MIMIC-IV. "
    "Available at: https://physionet.org/ (Accessed: 17 March 2026).",

    "Reimers, N. and Gurevych, I. (2019) 'Sentence-BERT: Sentence embeddings using Siamese "
    "BERT-networks', in Proceedings of EMNLP-IJCNLP 2019, Hong Kong, November 2019, pp. 3982–3992.",

    "Rezaei, M. R., Saadati Fard, R., Parker, J. L., Krishnan, R. G. and Lankarany, M. (2025) "
    "'Agentic medical knowledge graph RAG (AMGRAG): Dynamic PubMed-updated knowledge graph for "
    "medical QA', arXiv preprint arXiv:2503.00456.",

    "Rivera, J., Martinez, A., Lopez, R., Sanchez, D. and Perez, G. (2025) 'Real-world validation "
    "of MedSearch: RAG-based clinical search assistance', medRxiv, 2025.05.02.25320.",

    "Robertson, S. and Zaragoza, H. (2009) 'The probabilistic relevance framework: BM25 and beyond', "
    "Foundations and Trends in Information Retrieval, 3(4), pp. 333–389.",

    "Sharma, P., Kumar, A., Singh, R. and Patel, S. (2025) 'KGRAG: Knowledge graph-extended RAG "
    "for biomedical question answering', Applied Intelligence, 55, p. 1102.",

    "Singhal, K., Azizi, S., Tu, T., Mahdavi, S. S., Wei, J., Chung, H. W., Scales, N., Tanwani, A., "
    "Cole-Lewis, H., Pfohl, S., Payne, P., Seneviratne, M., Gamble, P., Kelly, C., Scharli, N., "
    "Markatou, M., Bhatt, H., Cole, J., Gottweis, J., Srinivasan, J., Miao, T., Nweke, E. C., "
    "Yakob, M., Parsons, T., Nori, H., King, E., Soltau, H., Macherey, K. and Dean, J. (2023) "
    "'Large language models encode clinical knowledge', Nature, 620, pp. 172–180.",

    "Vladika, J., Domres, A., Nguyen, M., Moser, R., Nano, J. and Matthes, F. (2025) 'Improving "
    "reliability and explainability of medical question answering with evidence-grounded chain-of-"
    "thought', arXiv preprint arXiv:2501.09876.",

    "World Health Organization (WHO) (2021) Ethics and governance of artificial intelligence for "
    "health: WHO guidance. Geneva: World Health Organization. "
    "Available at: https://www.who.int/publications/i/item/9789240029200 (Accessed: 17 March 2026).",

    "Xiong, G., Jin, Q., Lu, Z. and Zhang, A. (2024) 'Benchmarking retrieval-augmented generation "
    "for medicine', in Findings of the Association for Computational Linguistics: ACL 2024 "
    "(MIRAGE), Bangkok, August 2024.",

    "Yao, X., Wang, T., Zhang, S., Ma, Z., Wang, Y. and Liu, F. (2025) 'RAG-squared: Rationale-guided "
    "retrieval-augmented generation for medical QA', in Proceedings of ACL 2025.",

    "Zemchyk, A., Rao, S., Li, Y. and Mark, R. G. (2025) 'ArchEHR-QA 2025: Contrastive fine-tuning "
    "for retrieval-augmented biomedical QA', in Proceedings of BioNLP 2025.",
]
for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.first_line_indent = Inches(-0.4)
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(ref)
    set_font(run, size=11)

# ══════════════════════════════════════════════════════════════════════════════
#  SAVE
# ══════════════════════════════════════════════════════════════════════════════
out = "Research_Proposal_v4.docx"
doc.save(out)
print(f"Saved: {out}")
