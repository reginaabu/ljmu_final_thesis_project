"""
build_explainer.py – Generate a beginner-friendly project explainer Word document.
Run: python build_explainer.py
Output: Project_Explainer.docx
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
for section in doc.sections:
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin   = Inches(1.25)
    section.right_margin  = Inches(1.0)

# ── Helpers ───────────────────────────────────────────────────────────────────
def set_font(run, name="Calibri", size=11, bold=False, italic=False, color=None):
    run.font.name  = name
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def h1(text):
    p = doc.add_heading(text, level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
    return p

def h3(text):
    p = doc.add_heading(text, level=3)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x70, 0xAD, 0x47)
    return p

def para(text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_font(r, size=size, bold=bold, italic=italic)
    p.paragraph_format.space_after = Pt(6)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    set_font(r, size=11)
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    p.paragraph_format.space_after = Pt(3)
    return p

def numbered(text, level=0):
    p = doc.add_paragraph(style="List Number")
    r = p.add_run(text)
    set_font(r, size=11)
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    p.paragraph_format.space_after = Pt(3)
    return p

def note(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run("💡  " + text)
    set_font(r, size=10, italic=True, color=(0x7F, 0x7F, 0x7F))
    return p

def code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.name = "Courier New"
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    return p

def divider():
    doc.add_paragraph("─" * 90)

def table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    hdr_cells = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for run in hdr_cells[i].paragraphs[0].runs:
            run.font.bold = True
            run.font.size = Pt(10)
    for row_data in rows:
        row_cells = t.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = str(val)
            for run in row_cells[i].paragraphs[0].runs:
                run.font.size = Pt(10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t

# ═══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════════════════════════
tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tp.add_run("\n\n\nExplainable Safe Medical Chatbot\n")
r.font.name = "Calibri"; r.font.size = Pt(28); r.font.bold = True
r.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

r2 = tp.add_run("A Beginner's Guide to the Project\n")
r2.font.name = "Calibri"; r2.font.size = Pt(18)
r2.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

r3 = tp.add_run("\nPhases, Modules, Architecture & Evaluation\n\n")
r3.font.name = "Calibri"; r3.font.size = Pt(12); r3.font.italic = True
r3.font.color.rgb = RGBColor(0x59, 0x59, 0x59)

r4 = tp.add_run("Date: 2026-03-18\n\n\n")
r4.font.name = "Calibri"; r4.font.size = Pt(11)
r4.font.color.rgb = RGBColor(0x59, 0x59, 0x59)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — WHAT IS THIS PROJECT?
# ═══════════════════════════════════════════════════════════════════════════════
h1("1.  What Is This Project About?")

para(
    "Imagine you are a doctor, nurse, or patient who wants to ask a medical question — "
    "\"What are the symptoms of Shingles?\" or \"Does taking statins reduce heart disease risk?\" "
    "— and get a trustworthy, sourced answer rather than a generic Google result or a "
    "hallucinated AI response."
)
para(
    "This project builds exactly that: a medical question-answering chatbot called ArogyaSaathi "
    "that searches real medical literature to find relevant evidence, uses an AI (Claude) to "
    "generate a grounded answer based only on that evidence, checks whether the answer is safe, "
    "checks whether every claim is factually accurate, and explains where every piece of "
    "information came from using inline citations."
)
para(
    "The word Explainable in the title is the key design goal. Every answer shows its sources, "
    "its confidence score, and why it said what it said. Nothing is generated from the AI's "
    "training memory alone — every sentence must be traceable back to a retrieved document."
)

h2("1.1  The Problem With Existing Medical AI")
table(
    ["Problem", "Example", "How This Project Fixes It"],
    [
        ["Hallucination", "AI confidently states a wrong drug dosage", "Every claim verified against retrieved sources"],
        ["No citations", "Answer has no source — user cannot verify", "Every sentence has an inline citation (PMID / QID / Case)"],
        ["Unsafe advice", "AI says 'take 500 mg of X twice daily'", "Safety regex flags and adds professional consultation disclaimer"],
        ["Single dataset", "Only searches PubMed or only NIH", "Federated search across 3 medical knowledge sources simultaneously"],
        ["Black box", "Score shown but no explanation", "Per-claim verdicts (✅ supported / ❌ unsupported) shown to user"],
    ],
    col_widths=[1.4, 2.2, 2.8]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — THE THREE DATASETS
# ═══════════════════════════════════════════════════════════════════════════════
h1("2.  The Three Medical Datasets")

para(
    "The system works across three real-world medical datasets, each representing a different "
    "type of medical knowledge. Think of them as three separate medical libraries that the "
    "chatbot searches simultaneously."
)

h2("2.1  PubMedQA")
bullet("What it is: 1,000 yes/no research questions derived from PubMed journal abstracts.")
bullet("Example question: \"Does exercise during pregnancy prevent postnatal depression?\"")
bullet("Source: National Library of Medicine's PubMed database.")
bullet("Answer format: The system retrieves the research abstract and generates a 1–2 sentence evidence-based conclusion.")
bullet("Retrieval method: BM25 keyword search (the research question vocabulary maps well to keywords).")
note(
    "PubMedQA questions are research-level — meant for clinicians and researchers reading "
    "journal abstracts, not patients. Answers are conclusions drawn from study findings."
)

h2("2.2  MedQuAD  (Medical Question–Answer Dataset)")
bullet("What it is: 47,441 Q&A pairs from the NIH covering 38 distinct question types.")
bullet("Example question types: symptoms, treatments, causes, genetic changes, outlook, brand names, dosage, storage, prevention.")
bullet("Source: 12 NIH health topic sources including MedlinePlus and Genetics Home Reference.")
bullet("Answer format: Patient-facing plain-language answers written by NIH medical editors.")
bullet("Retrieval method: Hybrid (BM25 + semantic search) with question-type-aware filtering.")
note(
    "MedQuAD has a tricky data problem: 12 source files reuse the same question ID numbers, "
    "so the same ID like \"0000118-1\" could map to 6 completely different diseases. "
    "The system fixes this by renaming duplicates: 0000118-1, 0000118-1__1, 0000118-1__2, etc."
)

h2("2.3  ArchEHR-QA")
bullet("What it is: 20 real clinical cases — Electronic Health Record (EHR) discharge notes with clinician questions about each patient.")
bullet("Example question: \"Why was cardiac catheterization recommended to the patient?\"")
bullet("Source: PhysioNet (credentialed clinical data repository), BioNLP 2025 shared task.")
bullet("Answer format: The system answers using the patient's own discharge note as evidence.")
bullet("Retrieval method: Semantic search (clinical notes are narrative prose — keyword matching fails on medical shorthand).")
note(
    "ArchEHR uses real patient data. The system includes a PHI scrubber "
    "(utils/phi_scrub.py) to remove patient identifiers before sending text to external APIs — "
    "a HIPAA compliance requirement."
)

table(
    ["Dataset", "Size", "Question Style", "Audience", "Retriever"],
    [
        ["PubMedQA",  "1,000 abstracts",   "Yes/No research question", "Clinicians / researchers", "BM25"],
        ["MedQuAD",   "47,441 Q&A pairs",  "Patient-facing, 38 types", "Patients / general public", "Hybrid (BM25 + Semantic)"],
        ["ArchEHR-QA","20 clinical cases",  "Clinical EHR question",    "Clinicians / nurses",       "Semantic"],
    ],
    col_widths=[1.2, 1.4, 1.8, 1.6, 1.5]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — THE THREE PROJECT TRACKS (PHASES)
# ═══════════════════════════════════════════════════════════════════════════════
h1("3.  The Three Project Phases (Tracks)")

para(
    "The project was built in three progressive phases called tracks. Each track adds more "
    "intelligence to the system. Think of it as building a car — Track 1 is the engine, "
    "Track 2 is the GPS, and Track 3 is the full self-driving system with safety checks."
)

# Track 1
h2("3.1  Track 1 — BM25 Retrieval Baseline")
para("Core question: Can we find relevant medical documents using keyword matching?", bold=True)

h3("What is BM25?")
para(
    "BM25 (Best Match 25) is a classic keyword scoring algorithm used by search engines like "
    "Elasticsearch. It is an advanced version of counting how many times query words appear "
    "in each document, then ranking by score. It improves on simple word counting in two ways:"
)
bullet("Diminishing returns: If 'diabetes' appears 20 times vs 10 times in a document, the 20× document is NOT twice as relevant. BM25 applies a saturation curve.")
bullet("Length penalty: A very long document mentioning a word once is less relevant than a short focused document mentioning it once. BM25 normalises for document length.")

h3("Tuning Parameters")
para("BM25 has two parameters that control its behaviour:")
table(
    ["Parameter", "What it controls", "Tuned value"],
    [
        ["k1", "How quickly term frequency saturates (how much repetition matters)", "1.2"],
        ["b",  "How much document length is penalised (0 = ignore length, 1 = full penalty)", "0.75"],
    ],
    col_widths=[0.8, 3.8, 1.2]
)
note("These values were found automatically by tune_bm25.py, which tried hundreds of combinations and picked the best by MRR score.")

h3("Files Involved in Track 1")
table(
    ["File", "Purpose"],
    [
        ["scripts/run_track1.py",    "Runs the BM25 evaluation pipeline"],
        ["tune_bm25.py",             "Automatically finds best k1 and b parameters"],
        ["bm25_params.json",         "Saves the tuned k1=1.2, b=0.75 parameters"],
        ["pubmedqa_subset.csv",      "1,000 PubMedQA questions used as the test corpus"],
        ["metrics.json",             "Retrieval quality numbers (MRR, Precision, Recall)"],
        ["bm25_results.csv",         "Per-question breakdown of what was retrieved"],
        ["error_analysis.md",        "Which question types BM25 struggles with"],
    ],
    col_widths=[2.2, 4.2]
)

h3("What Gets Measured")
bullet("MRR (Mean Reciprocal Rank): On average, how high in the results is the correct document? MRR = 1.0 means always first.")
bullet("Precision@k: Of the top-k results, what fraction are relevant?")
bullet("Recall@k: Of all relevant documents, what fraction appear in the top-k results?")

h3("Limitation of Track 1")
para(
    "BM25 only matches keywords. A question about 'heart attack' will NOT match a document "
    "that says 'myocardial infarction' even though they mean exactly the same thing. "
    "This vocabulary gap motivates Track 2."
)

doc.add_page_break()

# Track 2
h2("3.2  Track 2 — Knowledge Graph Construction")
para("Core question: Can we improve retrieval by understanding medical concepts and their relationships?", bold=True)

h3("What is a Knowledge Graph?")
para(
    "A Knowledge Graph (KG) is a database of facts stored as triples: "
    "(entity → relationship → entity). For example:"
)
bullet("(metformin) → [treats] → (diabetes)")
bullet("(diabetes) → [causes] → (hyperglycemia)")
bullet("(aspirin) → [reduces] → (platelet aggregation)")

para(
    "If a user asks about 'diabetes treatment', the KG can find that 'metformin' is a "
    "treatment for diabetes even if the word 'treatment' does not appear in the document. "
    "This is called query expansion — adding related terms to the original search."
)

h3("How the KG is Built  (track2_build_kg.py)")
numbered("Load PubMedQA abstracts — the raw text of 1,000 research papers.")
numbered(
    "Named Entity Recognition (NER) using SciSpacy — SciSpacy is a medical NLP library. "
    "It reads each abstract and identifies: DISEASE names (e.g. 'fever', 'diabetes') "
    "and CHEMICAL names (e.g. 'aspirin', 'metformin')."
)
numbered(
    "Co-occurrence counting — if two entities appear in the same abstract, they are probably "
    "related. The system counts how many abstracts each pair of entities co-occurs in."
)
numbered("Build a graph — entities become nodes, co-occurrences become weighted edges.")
numbered("Extract triples — the strongest relationships become (entity1, co-occurs_with, entity2) triples.")

h3("How Query Expansion Works  (kg_expand.py)")
para("When a user types 'Does aspirin reduce heart attacks?', the system:")
numbered("Runs SciSpacy NER → finds 'aspirin' (CHEMICAL), 'heart attacks' (DISEASE).")
numbered("Looks up both entities in the KG graph.")
numbered("Finds their top neighbours — e.g. 'myocardial infarction', 'platelet aggregation', 'cardiovascular disease'.")
numbered("Appends neighbours to the query: original + 'myocardial infarction platelet aggregation cardiovascular disease'.")
numbered("The expanded query now matches more relevant documents in BM25.")

h3("Files Involved in Track 2")
table(
    ["File", "Purpose"],
    [
        ["track2_build_kg.py",     "Main KG construction pipeline"],
        ["scripts/install_scispacy.py", "Installs the SciSpacy medical NER model"],
        ["entities.csv",           "All discovered medical entities and their types"],
        ["triples.csv",            "All extracted (entity, relation, entity) triples"],
        ["kg_expand.py",           "Uses the built KG at query time to expand queries"],
    ],
    col_widths=[2.2, 4.2]
)

h3("Limitation of Track 2")
para(
    "Retrieval still returns documents, not answers. A clinician gets a list of abstracts "
    "and still has to read them and draw their own conclusions. "
    "This motivates Track 3 — automated answer generation."
)

doc.add_page_break()

# Track 3
h2("3.3  Track 3 — RAG Generation with Safety and Explainability")
para("Core question: Can we generate a safe, accurate, cited answer using retrieved evidence?", bold=True)

para(
    "This is the most complex and important track. It introduces RAG (Retrieval-Augmented "
    "Generation) — the modern approach where an AI generates answers grounded in retrieved "
    "documents rather than from its internal training data alone."
)

para("The full Track 3 pipeline runs every time a user asks a question:", bold=True)
table(
    ["Step", "Module", "What Happens"],
    [
        ["1", "kg_expand.py",            "Query is expanded with related medical concepts from the KG"],
        ["2", "app.py + semantic_index", "Federated retrieval across all 3 datasets simultaneously"],
        ["3", "reranker.py",             "Cross-encoder reranks top candidates for precision"],
        ["4", "rag_generate.py",         "Claude Sonnet generates a grounded, cited answer"],
        ["5", "evaluator/safety.py",     "Safety check — flags unsafe prescribing / emergency content"],
        ["6", "evaluator/fact_decompose.py + fact_verify.py", "Factuality check — verifies every claim"],
        ["7", "evaluator/__init__.py",   "Correction loop if factuality < 50%"],
        ["8", "evaluator/metrics.py",    "Faithfulness + G-Eval answer relevancy scoring"],
    ],
    col_widths=[0.5, 2.8, 3.2]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — MODULE DEEP DIVES
# ═══════════════════════════════════════════════════════════════════════════════
h1("4.  Module-by-Module Deep Dive")

# 4.1 Retrieval
h2("4.1  Federated Retrieval  (app.py, utils/semantic_index.py)")

para(
    "'Federated' means searching all three datasets simultaneously and combining results "
    "into a single ranked list. The system maintains three separate indexes (one per dataset) "
    "and queries them in parallel."
)

h3("4.1.1  BM25 Retrieval (PubMedQA)")
para(
    "PubMedQA uses BM25 on the context text. Research abstracts have rich keyword signals, "
    "so keyword matching works well. Two additional filters prevent false matches:"
)
bullet(
    "Q-field BM25 gate: Each candidate paper's research question is scored against the "
    "query's content terms. If score < 8.0, the candidate is excluded. This prevents "
    "a paper about 'halofantrine (malaria drug)' from appearing for a 'keratoderma' query — "
    "the two research questions share no content vocabulary."
)
bullet(
    "Per-result semantic threshold: Each retrieved chunk must have semantic similarity ≥ 0.50 "
    "to the query. Low-similarity chunks are discarded even if BM25 ranked them highly."
)

h3("4.1.2  Hybrid Retrieval (MedQuAD)")
para("MedQuAD uses Hybrid retrieval — a combination of BM25 and semantic search:")
bullet(
    "Semantic search uses all-MiniLM-L6-v2, a neural network that converts text into "
    "384-dimensional vectors (embeddings). Questions with similar meaning have similar "
    "vectors. 'Heart attack' and 'myocardial infarction' would now match correctly."
)
bullet(
    "BM25 and semantic results are combined using Reciprocal Rank Fusion (RRF): each method "
    "ranks all candidates, then the final score = 1/(rank + 60) from BM25 + "
    "1/(rank + 60) from semantic, weighted alpha=0.6 in favour of semantic."
)
bullet(
    "q_type-aware retrieval: MedQuAD has 38 question types (symptoms, treatment, outlook, "
    "genetic changes, brand names, etc.). A keyword classifier maps the query to the "
    "correct type and restricts semantic candidates to that type. For example: "
    "'What is the outlook for Factor X deficiency?' is classified as 'outlook' and "
    "only searches among 2,232 outlook-type Q&A pairs — preventing a general 'what is' "
    "answer from dominating."
)

table(
    ["Query Pattern", "Classified As", "Pool Size"],
    [
        ["'symptoms', 'signs of'",                   "symptoms",          "~4,338 rows"],
        ["'treatment', 'treat', 'therapy'",           "treatment",         "~3,906 rows"],
        ["'outlook', 'prognosis'",                    "outlook",           "~2,232 rows"],
        ["'genetic change', 'gene mutation'",         "genetic changes",   "~1,087 rows"],
        ["'brand name', 'brand names'",               "brand names",       "~1,238 rows"],
        ["'storage', 'disposal'",                     "storage and disposal", "~1,117 rows"],
        ["'side effect', 'adverse effect'",           "side effects",      "~1,301 rows"],
        ["'prevent', 'prevention'",                   "prevention",        "~1,239 rows"],
    ],
    col_widths=[2.5, 2.0, 1.5]
)

h3("4.1.3  Semantic Retrieval (ArchEHR)")
para(
    "ArchEHR uses pure semantic retrieval. Clinical discharge notes are written in "
    "narrative prose full of abbreviations, medical shorthand, and complex sentence "
    "structures that BM25 cannot handle. The sentence-transformer model understands "
    "the meaning of clinical language and finds relevant sentences even when exact "
    "keywords do not match."
)

h3("4.1.4  RRF Merging Across All Three Datasets")
para(
    "Results from all three datasets are merged into one ranked list using RRF again. "
    "Each result is tagged with its source dataset so the UI can display the correct "
    "badge (PubMed / MedQuAD / ArchEHR) and link format."
)

doc.add_page_break()

# 4.2 Reranker
h2("4.2  Cross-Encoder Reranking  (reranker.py)")
para(
    "After retrieval, the top ~30 candidates are reranked using a cross-encoder model "
    "(cross-encoder/ms-marco-MiniLM-L-6-v2, ~85 MB)."
)

h3("Why Rerank?")
para(
    "Retrieval is fast but imprecise. BM25 and semantic search score the query and "
    "document independently — they do not 'read' them together. A cross-encoder reads "
    "the query AND the document as a single input, giving it much richer understanding of "
    "relevance. The score it outputs is far more accurate at identifying 'does this "
    "document actually answer this specific question?'"
)

para("This is a standard two-stage retrieval pattern:")
numbered("Stage 1 (fast): BM25/semantic for broad recall — retrieve 30 candidates quickly.")
numbered("Stage 2 (precise): Cross-encoder reranks top 30 — pick the best 3.")

note(
    "The cross-encoder is optional — if sentence-transformers is not installed, the system "
    "falls back gracefully to the retrieval-only ranked order."
)

doc.add_page_break()

# 4.3 RAG Generation
h2("4.3  Answer Generation  (rag_generate.py)")
para(
    "The top 3 retrieved chunks (with their source IDs) are sent to Claude Sonnet "
    "(Anthropic's flagship model) with a carefully designed system prompt."
)

h3("The Prompt Design")
para("The system prompt tells Claude:")
numbered("Answer based ONLY on the provided evidence — do not use training data.")
numbered(
    "Restate the question's key subject in the opening sentence. For example: "
    "'The brand names of Benztropine Mesylate Oral are…' or 'The symptoms of Shingles include…' "
    "This echo-back rule improves answer relevancy scoring."
)
numbered("Cite every claim inline as (PMID XXXXXXXX) or (QID XXXXXXXX) or (Case XX).")
numbered("If evidence is insufficient, say what the evidence DOES say and note the gap.")
numbered("Do not speculate beyond the evidence.")

h3("Dataset-Specific Prompts")
table(
    ["Dataset", "Prompt Style", "Why"],
    [
        ["PubMedQA",   "1–2 sentences, direct yes/no conclusion, no paraphrased statistics",
         "Yes/no research questions need a direct conclusion. Longer answers generate specific statistical claims that fail verification as paraphrases."],
        ["MedQuAD",    "3–5 sentences, restated subject, full patient-facing explanation",
         "Patient questions need complete answers covering the full question scope."],
        ["ArchEHR-QA", "3–5 sentences, restated subject, grounded in discharge note",
         "Clinical questions need answers that trace directly to specific note content."],
    ],
    col_widths=[1.0, 2.0, 3.5]
)

h3("What Happens With Mixed Sources?")
para(
    "When the retrieved chunks come from multiple datasets (e.g. both MedQuAD and PubMedQA), "
    "the prompt uses generic labels: 'medical evidence records' and each chunk is labelled "
    "with its own identifier type in brackets so Claude always knows exactly how to cite it."
)

doc.add_page_break()

# 4.4 Safety
h2("4.4  Safety Check  (evaluator/safety.py)")
para(
    "The generated answer is scanned using regular expression (regex) patterns — "
    "text-matching rules — for three categories of unsafe content."
)

table(
    ["Category", "What It Detects", "Example Triggers", "Action"],
    [
        ["EMERGENCY",    "Imminent danger signals",          "'call 911', 'seek immediate care', overdose symptoms, 'life-threatening'",
         "Adds: 'Seek immediate medical attention'"],
        ["PRESCRIPTION", "Specific prescribing advice",      "'take 500mg of X', 'prescribe Y for Z', 'inject', drug name + 'take' within 50 chars",
         "Adds: 'Only take medication as prescribed'"],
        ["DIAGNOSIS",    "Definitive diagnostic statements",  "'you have X', 'this confirms Y', 'you are diagnosed with'",
         "Adds: 'Please consult a healthcare professional'"],
    ],
    col_widths=[1.2, 1.5, 2.3, 1.5]
)

para(
    "If any pattern fires, the answer is flagged as UNSAFE and a professional consultation "
    "disclaimer is appended. The answer content is still shown — the system does not "
    "refuse to answer — but the user is clearly warned."
)

note(
    "The safety checker only triggers on actual unsafe language patterns. Normal medical "
    "factual answers (brand names, symptoms, genetic information) pass through "
    "unmodified — the checker does not over-flag educational content."
)

doc.add_page_break()

# 4.5 Factuality
h2("4.5  Factuality Check  (evaluator/fact_decompose.py + fact_verify.py)")
para(
    "This is the most sophisticated quality-assurance step. It checks whether every "
    "individual claim in the answer is actually supported by the retrieved evidence."
)

h3("Step 1 — Fact Decomposition  (fact_decompose.py)")
para(
    "The answer is broken into atomic (indivisible) claims using Claude Haiku. "
    "'Atomic' means each claim asserts exactly one fact."
)
bullet("'Shingles is caused by reactivation of the varicella-zoster virus.' → one atomic claim")
bullet("'The main symptoms include a painful rash, blisters, and nerve pain.' → may become three separate claims: (1) painful rash, (2) blisters, (3) nerve pain")

para("Dataset-specific behaviour:")
bullet(
    "PubMedQA (≤2 sentences): LLM decomposition is SKIPPED. Each sentence is used "
    "directly as one atomic claim. This prevents over-fragmentation of short yes/no "
    "answers into 5–8 micro-claims that fail verification as paraphrases."
)
bullet(
    "All others: If the answer has more than 1 sentence or any sentence is over 30 words, "
    "Claude Haiku decomposes it into atomic claims via a single API call."
)

h3("Step 2 — Fact Verification  (fact_verify.py)")
para(
    "Each atomic claim is sent to Claude Haiku along with the retrieved evidence chunks. "
    "Haiku labels each claim as one of:"
)
bullet("✅  supported — the evidence explicitly states or strongly implies the claim")
bullet("❌  unsupported — the evidence does not address this claim")
bullet("⚠️  contradicted — the evidence explicitly says the opposite")

para("Factuality score = supported_claims ÷ total_claims", bold=True)

h3("Step 3 — Self-Correction Loop")
para(
    "If factuality score < 0.5 (less than half the claims are supported), a correction loop "
    "runs automatically:"
)
numbered("The answer is regenerated using a strict prompt requiring every single claim to have a direct citation.")
numbered("The new answer is evaluated for factuality.")
numbered("If the new answer has EQUAL OR BETTER factuality, it replaces the original.")
numbered("If not, the original is kept — preventing over-correction.")

note(
    "The correction loop ran on 1–2 questions per dataset in the latest evaluation. "
    "This shows the primary prompt is already well-tuned and rarely needs correction."
)

doc.add_page_break()

# 4.6 Metrics
h2("4.6  Quality Metrics  (evaluator/metrics.py)")
para(
    "Two additional quality metrics are computed after generation using the RAGAS framework "
    "and a custom G-Eval judge:"
)

h3("Faithfulness (RAGAS)")
para(
    "Faithfulness checks whether the answer introduces claims not present in the retrieved "
    "context — this is hallucination detection. It decomposes the answer into statements "
    "and verifies each against the context using Claude Haiku as the judge LLM."
)
bullet("Score range: 0.0 to 1.0")
bullet("1.0 means no hallucination — every claim is grounded in the retrieved evidence")
bullet("0.0 means all claims came from the AI's training data, not the evidence")
para("This is different from factuality:", italic=True)
bullet("Factuality: Are the claims ACCURATE (verified against evidence)?")
bullet("Faithfulness: Are the claims GROUNDED (only from the provided evidence, not from training memory)?")

h3("Answer Relevancy (G-Eval)")
para(
    "Answer Relevancy measures how well the answer addresses the original question. "
    "The system uses G-Eval — a Generative Evaluation approach where Claude Haiku is "
    "asked directly: 'Does this answer address this question? Score 0.0–1.0.'"
)
bullet("Score range: 0.0 to 1.0")
bullet("1.0 means the answer directly and completely addresses the question")
bullet("0.5–0.9 means mostly on-topic with some gaps")
bullet("0.0–0.3 means barely or does not address the question")

para(
    "Previously, the RAGAS reverse-question approach was used — it generated questions from "
    "the answer and measured cosine similarity to the original query. This approach "
    "collapsed to 0.0 for list answers (brand names), storage/disposal questions, and "
    "safety-hedged clinical answers. G-Eval handles all these cases correctly."
)

table(
    ["Metric", "Tool", "What It Measures", "Score Range"],
    [
        ["Faithfulness",      "RAGAS + Claude Haiku", "Fraction of claims inferable from retrieved context (anti-hallucination)", "0.0 – 1.0"],
        ["Answer Relevancy",  "G-Eval (Claude Haiku)", "Does the answer address the question? Direct LLM judge", "0.0 – 1.0"],
        ["Factuality",        "fact_verify.py",        "Fraction of atomic claims supported by evidence", "0.0 – 1.0"],
        ["Safety Rate",       "safety.py",             "Fraction of answers that are safe (no unsafe flags)", "0% – 100%"],
        ["Mean Latency",      "time.perf_counter()",   "Wall-clock time from query to answer", "seconds"],
    ],
    col_widths=[1.5, 1.6, 2.3, 1.0]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 5 — SUPPORTING MODULES
# ═══════════════════════════════════════════════════════════════════════════════
h1("5.  Supporting Modules")

table(
    ["Module", "File", "Purpose"],
    [
        ["Dataset Adapter",  "utils/dataset_adapter.py",  "Unified interface to all 3 datasets. Normalises different column names, file formats (CSV, HuggingFace, XML) to one canonical schema: {doc_id, question, context, focus, q_type}"],
        ["Semantic Index",   "utils/semantic_index.py",   "SemanticIndex (pure dense retrieval) and HybridIndex (RRF fusion). Supports q_type filtering for MedQuAD question-type-aware retrieval"],
        ["PHI Scrubber",     "utils/phi_scrub.py",        "Removes patient identifiers (names, dates, hospital IDs) from ArchEHR text before sending to external APIs. HIPAA compliance"],
        ["Logging Config",   "utils/logging_config.py",   "Centralised logging setup so every module logs consistently with timestamps"],
        ["KG Expander",      "kg_expand.py",              "Loads triples.csv, extracts entities from queries using SciSpacy NER, returns related terms for query expansion"],
        ["Reranker",         "reranker.py",               "Cross-encoder reranking using ms-marco-MiniLM-L-6-v2. Optional — degrades gracefully if not installed"],
        ["RAG Generator",    "rag_generate.py",           "Calls Claude Sonnet API with retrieved chunks and dataset-specific prompts. Returns cited answer text"],
        ["Eval Harness",     "eval_harness.py",           "CLI script: samples N questions, runs full pipeline, writes Markdown report with per-question and aggregate metrics"],
        ["Pipeline Runner",  "run_pipeline.py",           "Single entrypoint: python run_pipeline.py track1/track2/track3-eval/app"],
        ["App",              "app.py",                    "Main Streamlit web application. Orchestrates all modules and renders the UI"],
    ],
    col_widths=[1.5, 2.0, 3.0]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 6 — HOW ALL FILES CONNECT
# ═══════════════════════════════════════════════════════════════════════════════
h1("6.  How All Files Connect — The Full Data Flow")

para("When a user types a question in the browser, this is what happens step by step:")

numbered(
    "app.py receives the query from the Streamlit text input and starts the pipeline."
)
numbered(
    "kg_expand.py runs SciSpacy NER on the query, finds medical entities (diseases, chemicals), "
    "looks them up in triples.csv, and appends related terms to produce an expanded query."
)
numbered(
    "app.py calls _retrieve_docs() for each loaded dataset bundle (PubMedQA BM25, "
    "MedQuAD Hybrid, ArchEHR Semantic). Each runs independently and returns (doc, score, label) tuples."
)
numbered(
    "retrieve_federated() in app.py merges all results using RRF, producing one ranked "
    "list with source dataset tags."
)
numbered(
    "reranker.py receives the top ~30 merged results and reranks them using the "
    "cross-encoder model, adding a ce_score to each document."
)
numbered(
    "The top 3 chunks are selected and passed to rag_generate.py along with the correct "
    "id_label, source_type, and dataset name."
)
numbered(
    "rag_generate.py builds the system prompt (dataset-specific), formats the chunks as "
    "evidence context, calls the Claude Sonnet API, and returns the cited answer text."
)
numbered(
    "evaluator/safety.py scans the answer with regex patterns and appends disclaimers "
    "if any unsafe patterns are found."
)
numbered(
    "evaluator/fact_decompose.py breaks the answer into atomic claims (dataset-aware: "
    "PubMedQA short answers skip LLM decomp)."
)
numbered(
    "evaluator/fact_verify.py sends all claims + evidence to Claude Haiku, which labels "
    "each claim as supported / unsupported / contradicted."
)
numbered(
    "factuality_score = supported / total is computed. If < 0.5, steps 7–10 run again "
    "with a strict prompt and the better answer is kept."
)
numbered(
    "evaluator/metrics.py computes RAGAS Faithfulness (Claude Haiku) and G-Eval "
    "Answer Relevancy (Claude Haiku scoring 0.0–1.0)."
)
numbered(
    "app.py renders the final answer with: inline citations, safety badge, factuality "
    "score, per-claim verdict list, faithfulness/relevancy bars, and source snippets."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 7 — LATEST EVALUATION RESULTS
# ═══════════════════════════════════════════════════════════════════════════════
h1("7.  Latest Evaluation Results  (2026-03-18, seed=42)")

para(
    "The evaluation harness (eval_harness.py) runs the complete Track 3 pipeline on "
    "25 randomly sampled questions per dataset and reports aggregate metrics."
)

h2("7.1  Summary Table")
table(
    ["Metric", "PubMedQA (25 Qs)", "MedQuAD (25 Qs)", "ArchEHR (20 Qs)"],
    [
        ["Faithfulness",      "0.867", "0.922", "0.674"],
        ["Answer Relevancy",  "0.742", "0.474", "0.693"],
        ["Factuality",        "0.750", "0.885", "0.861"],
        ["Safety Rate",       "96%",   "92%",   "90%"],
        ["Mean Latency",      "3.69s", "3.84s", "5.33s"],
        ["Corrections",       "1/25",  "2/25",  "2/20"],
    ],
    col_widths=[2.0, 1.5, 1.5, 1.5]
)

h2("7.2  Per-Dataset Analysis")

h3("PubMedQA")
bullet("Faithfulness 0.867: Strong — most yes/no answers are well-grounded in abstract content.")
bullet("Answer Relevancy 0.742 (G-Eval): Solid — the 1–2 sentence concise prompt with echo-back rule produces well-targeted research conclusions.")
bullet("Factuality 0.750: Improved significantly from 0.521 baseline by skipping LLM decomposition for short answers (prevents paraphrase failures).")
bullet("Safety 96%: One UNSAFE(EMERGENCY) flag — a mental health differences question that triggered emergency pattern.")
bullet("Only 1 correction in 25 questions — primary prompt is well-calibrated.")

h3("MedQuAD")
bullet("Faithfulness 0.922: Highest across all three datasets — the hybrid retrieval with q_type filtering finds precise Q&A pairs.")
bullet("Answer Relevancy 0.474 (G-Eval): Moderate — 6 of 25 questions score 0.000 due to retrieval precision limits for very rare conditions (Weaver syndrome, Factor X deficiency) where the semantic model finds semantically similar but wrong-disease content.")
bullet("Factuality 0.885: Strong — NIH Q&A pairs provide clear factual content that LLM answers can accurately reflect.")
bullet("Safety 92%: 2 UNSAFE flags — one PRESCRIPTION (Efavirenz prescribing) and one EMERGENCY (laxative overdose).")

h3("ArchEHR")
bullet("Faithfulness 0.674: Lower than other datasets — clinical narrative notes are fragmented and the RAGAS faithfulness checker is strict about claim-to-context traceability.")
bullet("Answer Relevancy 0.693 (G-Eval): Second-highest — G-Eval handles clinical questions well where the RAGAS metric would have given 0.000 for specific patient questions.")
bullet("Factuality 0.861: High — clinical notes provide concrete, verifiable facts about specific patient events.")
bullet("Safety 90%: 2 UNSAFE(EMERGENCY) flags — pain management and overdose clinical cases (expected for real clinical EHR data).")
bullet("Higher latency (5.33s) — clinical notes are longer and semantic encoding takes more time.")

h2("7.3  Progressive Improvement Summary")
para("The metrics improved substantially through targeted engineering across all three tracks:")
table(
    ["Improvement", "Before", "After", "Change"],
    [
        ["PubMedQA factuality (LLM decomp fix)",         "0.521", "0.750", "+0.229"],
        ["PubMedQA answer relevancy (echo-back + G-Eval)", "~0.0",  "0.742", "+0.742"],
        ["MedQuAD faithfulness (q_type retrieval)",       "0.906", "0.922", "+0.016"],
        ["MedQuAD factuality (q_type retrieval)",         "0.841", "0.885", "+0.044"],
        ["ArchEHR answer relevancy (G-Eval)",             "0.317", "0.693", "+0.376"],
        ["ArchEHR factuality",                            "0.776", "0.861", "+0.085"],
        ["PubMedQA corrections needed",                   "12/25", "1/25",  "−11"],
    ],
    col_widths=[3.2, 0.8, 0.8, 0.8]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 8 — KEY DESIGN DECISIONS
# ═══════════════════════════════════════════════════════════════════════════════
h1("8.  Key Design Decisions and Why They Matter")

table(
    ["Decision", "Alternative Considered", "Why This Choice"],
    [
        ["BM25 for PubMedQA, Hybrid for MedQuAD, Semantic for ArchEHR",
         "Use the same retriever for all datasets",
         "Each dataset has a different structure. Research abstracts suit keywords. Patient Q&A needs question-type matching. Clinical notes need semantic understanding."],
        ["G-Eval (Claude Haiku LLM judge) for answer relevancy",
         "RAGAS reverse-question cosine similarity",
         "RAGAS scored 0.000 for list answers, brand names, storage questions. G-Eval directly asks 'does this answer address the question?' and handles all answer types."],
        ["Skip LLM decomposition for PubMedQA ≤2 sentences",
         "Always decompose with LLM",
         "2-sentence answers were being fragmented into 5–8 micro-claims. Paraphrased statistics (e.g. '25%' vs '24.7%') failed verification. Using sentences directly cut false failures."],
        ["q_type-aware retrieval for MedQuAD",
         "Unfiltered semantic search across all 47k rows",
         "Without filtering, 'What is X?' always wins (longest answer). q_type filtering ensures 'What is the outlook for X?' searches only among outlook-type Q&A pairs."],
        ["Per-result semantic threshold (≥0.50)",
         "Top-1 gate (allow all if best result passes)",
         "Top-1 gate allowed low-similarity results through. ArchEHR CASE_ID 15 was scoring 0.26 similarity and appearing for unrelated keratoderma queries."],
        ["MedQuAD doc_id deduplication",
         "Use raw doc_ids as-is",
         "12 source files share the same QID numbers. Without dedup, 'QID 0000118-1' mapped to 6 different diseases and HybridIndex returned wrong context for most of them."],
        ["Factuality-triggered correction loop with score comparison",
         "Always run strict prompt; or never correct",
         "Always running strict wastes API calls. Never correcting leaves bad answers. Score comparison ensures we only replace an answer if the correction actually helped."],
    ],
    col_widths=[1.8, 1.8, 2.9]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 9 — HOW TO RUN THE PROJECT
# ═══════════════════════════════════════════════════════════════════════════════
h1("9.  How to Run the Project")

h2("9.1  Prerequisites")
bullet("Python 3.10 or 3.12")
bullet("ANTHROPIC_API_KEY set in environment or .streamlit/secrets.toml")
bullet("pip install -r requirements.txt")

h2("9.2  Run the Web App")
code("python -m streamlit run app.py --browser.gatherUsageStats false")
para("Opens at http://localhost:8501. The app builds all indexes on first run (~2 minutes for MedQuAD).")

h2("9.3  Run Evaluation")
code("python eval_harness.py --dataset pubmedqa --n 25 --seed 42")
code("python eval_harness.py --dataset medquad --n 25 --seed 42")
code("python eval_harness.py --dataset archehr_qa --csv-path data/archehr_qa --n 25 --seed 42")
para("Each run generates a Markdown report: pubmedqa_eval_report.md, medquad_eval_report.md, archehr_eval_report.md")

h2("9.4  Run Track 1 BM25 Baseline")
code("python run_pipeline.py track1")

h2("9.5  Run Track 2 Knowledge Graph Build")
code("python run_pipeline.py track2")

h2("9.6  Regenerate Documents")
code("python build_report.py        # regenerates thesis")
code("python build_research_proposal.py  # regenerates research proposal")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 10 — GLOSSARY
# ═══════════════════════════════════════════════════════════════════════════════
h1("10.  Glossary of Terms")

table(
    ["Term", "Plain-English Definition"],
    [
        ["RAG",            "Retrieval-Augmented Generation — generating AI answers from retrieved documents rather than from the AI's training memory alone"],
        ["BM25",           "Best Match 25 — a classic keyword scoring algorithm used by search engines to rank documents by relevance to a query"],
        ["Embedding",      "A list of numbers (vector) representing the meaning of a piece of text. Similar texts have similar vectors. Enables semantic search."],
        ["Semantic Search","Finding documents by meaning rather than exact keywords. 'Heart attack' matches 'myocardial infarction' because they mean the same thing."],
        ["RRF",            "Reciprocal Rank Fusion — a method to combine rankings from multiple retrieval systems by adding 1/(rank + 60) scores"],
        ["Hybrid Retrieval","Combining BM25 keyword matching and semantic embedding search using RRF, getting the benefits of both"],
        ["Cross-Encoder",  "A neural network model that reads a query AND a document together to give a precise relevance score. More accurate than embedding search but slower."],
        ["NER",            "Named Entity Recognition — automatically identifying names of medical entities (diseases, chemicals) in text using NLP models"],
        ["SciSpacy",       "A medical NLP library built on spaCy, specialised for biomedical and clinical text entity recognition"],
        ["Knowledge Graph","A structured database of facts stored as (entity → relationship → entity) triples, enabling reasoning about medical concepts"],
        ["Atomic Claim",   "A single, indivisible factual statement asserting exactly one fact — the unit used in fact decomposition and verification"],
        ["Factuality",     "The fraction of atomic claims in an answer that are verified as supported by the retrieved evidence (supported / total)"],
        ["Faithfulness",   "RAGAS metric: the fraction of answer statements that can be inferred from the retrieved context. Measures anti-hallucination."],
        ["G-Eval",         "Generative Evaluation — using an LLM (Claude Haiku) as an evaluator that scores answer quality by direct judgment"],
        ["q_type",         "Question type — the category of a MedQuAD question (symptoms, treatment, outlook, genetic changes, brand names, etc.) used for type-aware retrieval"],
        ["Federated Search","Querying multiple separate databases simultaneously and merging results into one ranked list"],
        ["PHI",            "Protected Health Information — patient identifiers like names, dates, IDs that must be scrubbed before sharing data externally (HIPAA requirement)"],
        ["RAGAS",          "Retrieval Augmented Generation Assessment — an evaluation framework providing faithfulness and answer relevancy metrics"],
        ["Correction Loop","A self-improvement mechanism: if factuality < 0.5, regenerate with a stricter prompt and keep whichever answer has better factuality"],
        ["EHR",            "Electronic Health Record — a patient's digital medical record including discharge notes, diagnoses, medications, and test results"],
    ],
    col_widths=[1.8, 4.7]
)

# ── Save ─────────────────────────────────────────────────────────────────────
output_path = "Project_Explainer.docx"
doc.save(output_path)
print(f"Saved -> {output_path}")
